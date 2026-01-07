#!/usr/bin/env python3
"""Build a single Word document containing all exported figures.

Conservative behavior:
- Includes image files only (.png/.jpg/.jpeg).
- Uses filename as the figure label.

Usage:
  python3 3_Analysis/4_Plots_Code/build_all_figures_docx.py \
    --fig_dir 4_Model_Results/Figures \
    --out 4_Model_Results/Summary/All_Figures_20260106.docx
"""

from __future__ import annotations

import argparse
from pathlib import Path

from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH


IMAGE_EXTS = {".png", ".jpg", ".jpeg"}


def set_doc_font(doc: Document, name: str = "Times New Roman", size_pt: int = 12) -> None:
    style = doc.styles["Normal"]
    style.font.name = name
    style.font.size = Pt(size_pt)


def natural_key(p: Path) -> list[object]:
    # Natural-ish sort: splits digits for ordering fig2 before fig10.
    parts: list[object] = []
    s = p.name
    buf = ""
    is_digit = None
    for ch in s:
        if ch.isdigit():
            if is_digit is False:
                parts.append(buf.lower())
                buf = ""
            buf += ch
            is_digit = True
        else:
            if is_digit is True:
                parts.append(int(buf))
                buf = ""
            buf += ch
            is_digit = False
    if buf:
        parts.append(int(buf) if (is_digit is True) else buf.lower())
    return parts


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--fig_dir", required=True, help="Directory containing exported figure images")
    ap.add_argument("--out", required=True, help="Output .docx path")
    ap.add_argument("--width_inches", type=float, default=6.5, help="Image width in inches")
    args = ap.parse_args()

    fig_dir = Path(args.fig_dir)
    out_path = Path(args.out)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    images = [p for p in fig_dir.iterdir() if p.is_file() and p.suffix.lower() in IMAGE_EXTS]
    images.sort(key=natural_key)

    doc = Document()
    set_doc_font(doc)

    title = doc.add_paragraph("All Figures")
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title.runs[0].bold = True
    title.runs[0].font.size = Pt(14)

    subtitle = doc.add_paragraph(f"Source folder: {fig_dir.as_posix()}")
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    if not images:
        doc.add_paragraph("No image files were found in the specified folder.")
        doc.save(out_path.as_posix())
        return

    for idx, img in enumerate(images):
        doc.add_page_break()

        p = doc.add_paragraph(img.name)
        p.runs[0].bold = True
        p.paragraph_format.space_after = Pt(6)

        # Insert image
        doc.add_picture(img.as_posix(), width=Inches(args.width_inches))

        # Center the picture paragraph (last paragraph contains the picture)
        try:
            doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
        except Exception:
            pass

    doc.save(out_path.as_posix())


if __name__ == "__main__":
    main()
