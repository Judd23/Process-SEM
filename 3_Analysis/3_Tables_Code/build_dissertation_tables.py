#!/usr/bin/env python3
"""
Build Complete APA 7 Dissertation Tables for Process-SEM Analysis

Generates all required tables per JARS guidelines:
- Tables 1-3: Sample, Descriptives, Missing Data
- Tables 4-6: PSW Diagnostics (PS model, balance, weights)
- Tables 7-9: Measurement Model (CFA, invariance, fit)
- Tables 10-12: Structural Results (paths, indirect, conditional)
- Table 13: Robustness Checks

APA 7 Table Formatting (per Publication Manual, 7th ed.):
============================================================
Table Structure:
  - Table number: Bold, flush left (e.g., "Table 1")
  - Title: Italic, title case, flush left, double-spaced below number
  - Headings: 
    * Stub heading: leftmost column header (typically "Variable")
    * Column headers: single-column descriptive headers
    * Column spanners: headers spanning 2+ columns (decked heads)
    * Table spanners: headers spanning full table width
  - Body: Data rows with stub column left-aligned, numbers right-aligned

Borders (3 horizontal lines only):
  - Top border: above column spanners or headers
  - Header border: below column headers (separating head from body)
  - Bottom border: below last data row
  - No vertical lines, no internal horizontal lines

Notes (below table, flush left):
  - General note: Begins with "Note." in italics
  - Specific notes: Use superscript letters (ᵃ, ᵇ, ᶜ)
  - Probability notes: Asterisks for significance (*p < .05, **p < .01)

Usage:
    python scripts/build_dissertation_tables.py --outdir <path> [options]

Output: Dissertation_Tables.docx with all tables in APA 7 format
"""

import argparse
import json
from pathlib import Path
import pandas as pd
import numpy as np
from docx import Document
from docx.shared import Pt, Inches, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.section import WD_ORIENT
from docx.oxml.ns import qn, nsdecls
from docx.oxml import parse_xml


# =============================================================================
# FORMATTING UTILITIES
# =============================================================================

# Variable name mapping: technical -> readable
VARIABLE_LABELS = {
    # Treatment variables
    'x_FASt': 'FASt Status (≥12 credits)',
    'credit_dose': 'Dual Credit Units Earned',
    'credit_dose_c': 'Credit Dose (centered)',
    'XZ_c': 'FASt × Credit Dose',
    
    # Covariates
    '(Intercept)': 'Intercept',
    'cohort': 'Cohort (2022–23 vs 2023–24)',
    'hgrades': 'High School GPA',
    'hgrades_c': 'HS GPA (centered)',
    'bparented': 'Parent Education Level',
    'bparented_c': 'Parent Education (centered)',
    'pell': 'Pell Grant Recipient',
    'hapcl': 'AP/IB/College-Level Courses',
    'hprecalc13': 'Precalculus or Higher',
    'hchallenge': 'HS Academic Challenge',
    'hchallenge_c': 'HS Challenge (centered)',
    
    # Mediators - Emotional Distress
    'EmoDiss': 'Emotional Distress',
    'MHWdacad': 'Difficulty: Academic Work',
    'MHWdlonely': 'Difficulty: Loneliness',
    'MHWdmental': 'Difficulty: Mental Health',
    'MHWdexhaust': 'Difficulty: Exhaustion',
    'MHWdsleep': 'Difficulty: Sleep',
    'MHWdfinancial': 'Difficulty: Finances',
    
    # Mediators - Quality of Engagement
    'QualEngag': 'Quality of Engagement',
    'QIstudent': 'Quality: Student Interactions',
    'QIadvisor': 'Quality: Advisor Interactions',
    'QIfaculty': 'Quality: Faculty Interactions',
    'QIstaff': 'Quality: Staff Interactions',
    'QIadmin': 'Quality: Admin Interactions',
    
    # Outcome - Developmental Adjustment (2nd order)
    'DevAdj': 'Developmental Adjustment',
    
    # Belonging subscale
    'Belong': 'Sense of Belonging',
    'sbvalued': 'Belonging: Feel Valued',
    'sbmyself': 'Belonging: Can Be Myself',
    'sbcommunity': 'Belonging: Sense of Community',
    
    # Perceived Gains subscale
    'Gains': 'Perceived Gains',
    'pgthink': 'Gain: Critical Thinking',
    'pganalyze': 'Gain: Analyzing Problems',
    'pgwork': 'Gain: Working with Others',
    'pgvalues': 'Gain: Personal Values',
    'pgprobsolve': 'Gain: Problem Solving',
    
    # Supportive Environment subscale
    'SupportEnv': 'Supportive Environment',
    'SEwellness': 'Support: Physical Wellness',
    'SEnonacad': 'Support: Non-Academic',
    'SEactivities': 'Support: Activities',
    'SEacademic': 'Support: Academics',
    'SEdiverse': 'Support: Diverse Interactions',
    
    # Satisfaction subscale
    'Satisf': 'Satisfaction',
    'evalexp': 'Overall Experience Rating',
    'sameinst': 'Would Choose Same Institution',
    
    # Other study variables
    'SFcareer': 'Career Preparation',
    'SFotherwork': 'Other Work Activities',
    'SFdiscuss': 'Discussion Activities',
    'SFperform': 'Performance Activities',
    
    # Summary statistics (for balance table)
    'Mean |SMD|': 'Mean |SMD|',
    'Max |SMD|': 'Max |SMD|',
}


def get_label(var_name):
    """Get readable label for a variable name."""
    return VARIABLE_LABELS.get(var_name, var_name)


def find_csv(data_dir, filename):
    """Find CSV file checking both main directory and pooled subfolder."""
    # Try main directory first
    main_path = data_dir / filename
    if main_path.exists():
        return main_path
    # Try pooled subfolder
    pooled_path = data_dir / "pooled" / filename
    if pooled_path.exists():
        return pooled_path
    # Return main path (will be checked with exists() later)
    return main_path


def load_standardized_coefficients(data_dir):
    """Load standardized coefficients from CSV or TXT file.
    
    Returns dict mapping parameter labels to standardized coefficients (std.all or est.std).
    """
    std_map = {}
    
    # Try CSV format first (point_estimates_parameter_table_std.csv)
    csv_path = find_csv(data_dir, "point_estimates_parameter_table_std.csv")
    if csv_path.exists():
        std_df = pd.read_csv(csv_path)
        # Filter to regression paths and defined parameters with labels
        for _, row in std_df.iterrows():
            if row['op'] in ('~', ':=') and pd.notna(row.get('label')) and row.get('label') != '':
                std_col = 'std.all' if 'std.all' in std_df.columns else 'est.std'
                if std_col in std_df.columns:
                    std_map[row['label']] = row[std_col]
        return std_map
    
    # Try TXT format (structural_standardizedSolution.txt) - tab-delimited
    txt_path = data_dir / "structural_standardizedSolution.txt"
    if txt_path.exists():
        try:
            std_df = pd.read_csv(txt_path, sep='\t')
            # The TXT format has columns: lhs, op, rhs, label, est.std, se, z, pvalue, ci.lower, ci.upper
            for _, row in std_df.iterrows():
                if row['op'] in ('~', ':=') and pd.notna(row.get('label')) and str(row.get('label', '')).strip() != '':
                    std_col = 'est.std' if 'est.std' in std_df.columns else 'est'
                    if std_col in std_df.columns:
                        std_map[row['label']] = row[std_col]
        except Exception as e:
            print(f"Warning: Could not parse {txt_path}: {e}")
    
    return std_map


def fmt(x, nd=2):
    """Format number to nd decimal places."""
    try:
        v = float(x)
        if pd.isna(v):
            return "—"
        return f"{v:.{nd}f}"
    except:
        return str(x) if pd.notna(x) else "—"


def fmt_ci(lo, hi, nd=2):
    """Format confidence interval [lo, hi]."""
    return f"[{fmt(lo, nd)}, {fmt(hi, nd)}]"


def fmt_pct(x, nd=1):
    """Format as percentage."""
    try:
        return f"{float(x):.{nd}f}%"
    except:
        return "—"


def fmt_int(x):
    """Format as integer with comma separator."""
    try:
        return f"{int(x):,}"
    except:
        return str(x) if pd.notna(x) else "—"


def set_table_apa_borders(table):
    """Apply APA 7 borders: top, below header, bottom only."""
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else parse_xml(r'<w:tblPr %s/>' % nsdecls('w'))
    tblBorders = parse_xml(
        r'<w:tblBorders %s>'
        r'<w:top w:val="single" w:sz="4" w:color="000000"/>'
        r'<w:bottom w:val="single" w:sz="4" w:color="000000"/>'
        r'<w:insideH w:val="nil"/>'
        r'<w:insideV w:val="nil"/>'
        r'<w:left w:val="nil"/>'
        r'<w:right w:val="nil"/>'
        r'</w:tblBorders>' % nsdecls('w')
    )
    tblPr.append(tblBorders)


def set_row_border_bottom(row):
    """Add bottom border to a row (for header separation)."""
    for cell in row.cells:
        tc = cell._tc
        tcPr = tc.get_or_add_tcPr()
        tcBorders = parse_xml(
            r'<w:tcBorders %s>'
            r'<w:bottom w:val="single" w:sz="4" w:color="000000"/>'
            r'</w:tcBorders>' % nsdecls('w')
        )
        tcPr.append(tcBorders)


def set_cell_border_bottom(cell):
    """Add bottom border to a single cell (for spanner-length separators)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        r'<w:tcBorders %s>'
        r'<w:bottom w:val="single" w:sz="4" w:color="000000"/>'
        r'</w:tcBorders>' % nsdecls('w')
    )
    tcPr.append(tcBorders)


def set_cell_border_top(cell):
    """Add top border to a single cell (for total row separators)."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        r'<w:tcBorders %s>'
        r'<w:top w:val="single" w:sz="4" w:color="000000"/>'
        r'</w:tcBorders>' % nsdecls('w')
    )
    tcPr.append(tcBorders)


def merge_cells_horizontal(table, row_idx, start_col, end_col):
    """Merge cells horizontally for column spanners.
    
    Merges even single column (for consistent formatting).
    """
    if end_col < start_col:
        return  # Invalid range
    if end_col == start_col:
        return  # Single column, no merge needed but still valid
    row = table.rows[row_idx]
    start_cell = row.cells[start_col]
    for col in range(start_col + 1, end_col + 1):
        start_cell.merge(row.cells[col])


def set_cell_text(cell, text, bold=False, italic=False, align='center', font_size=11):
    """Set cell text with formatting.
    
    Args:
        cell: Table cell to format
        text: Text content
        bold: Bold text
        italic: Italic text
        align: 'left', 'center', or 'right'
        font_size: Font size in points
    """
    cell.text = str(text)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    if align == 'center':
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
    elif align == 'right':
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    else:
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT
    for run in cell.paragraphs[0].runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(font_size)
        run.bold = bold
        run.italic = italic


def add_apa7_table_advanced(doc, table_num, title, data_rows, 
                            stub_heading=None,
                            column_headers=None,
                            column_spanners=None,
                            table_spanner=None,
                            total_row_indices=None,
                            note_general=None, 
                            notes_specific=None, 
                            note_probability=None, 
                            compact=False):
    """
    Add APA 7 formatted table with full header structure.
    
    APA 7 Table Structure (stacked top to bottom):
    ─────────────────────────────────────────────── (top border)
    Table spanner: Header covering entire table width (optional)
    ─────────────────────────────────────────────── (full-width if table spanner)
    Column spanners: Headers spanning columns (decked heads)
           ───────────    ─────────── (category-length borders under each spanner)
    Column headers: Individual column headers  
    ─────────────────────────────────────────────── (header border)
    Data rows...
    ─────────────────────────────────────────────── (border above totals if specified)
    Total/summary rows...
    ─────────────────────────────────────────────── (bottom border)
    
    Args:
        doc: Document object
        table_num: Table number
        title: Table title (italic)
        data_rows: List of lists containing data
        stub_heading: Text for leftmost column header (e.g., "Variable", "Covariate")
        column_headers: List of column header texts
        column_spanners: List of (text, start_col, end_col) tuples for spanning headers
        table_spanner: Text for full-width header spanning all columns
        total_row_indices: List of row indices (0-based) that are totals (get top border)
        note_general: General note text (after "Note. ")
        notes_specific: List of specific notes (with superscripts)
        note_probability: Probability note (e.g., "*p < .05")
        compact: If True, reduce spacing
    """
    spacing = Pt(6) if compact else Pt(12)
    font_size = 10 if compact else 11
    note_size = 9 if compact else 10
    
    # Table number (bold)
    p_num = doc.add_paragraph()
    p_num.paragraph_format.space_before = spacing
    p_num.paragraph_format.space_after = Pt(0)
    run_num = p_num.add_run(f"Table {table_num}")
    run_num.bold = True
    run_num.font.name = 'Times New Roman'
    run_num.font.size = Pt(12)
    
    # Title (italic)
    p_title = doc.add_paragraph()
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(6)
    run_title = p_title.add_run(title)
    run_title.italic = True
    run_title.font.name = 'Times New Roman'
    run_title.font.size = Pt(12)
    
    if not data_rows or not column_headers:
        return table_num + 1
    
    # Calculate number of header rows needed
    num_header_rows = 1  # Always have column headers
    if column_spanners:
        num_header_rows = 2  # Add row for spanners (decked heads)
    if table_spanner:
        num_header_rows += 1  # Add row for table spanner
    
    num_cols = len(column_headers)
    num_data_rows = len(data_rows)
    
    # Create table
    table = doc.add_table(rows=num_header_rows + num_data_rows, cols=num_cols)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = False  # Disable autofit for explicit widths
    set_table_apa_borders(table)
    
    # Set column widths - first column wider for stub, others equal
    total_width = Inches(6.5)  # Standard page width minus margins
    stub_width = Inches(2.0)   # First column (variable names)
    other_width = (total_width - stub_width) / max(1, num_cols - 1)
    
    for row in table.rows:
        for col_idx, cell in enumerate(row.cells):
            if col_idx == 0:
                cell.width = stub_width
            else:
                cell.width = other_width
    
    current_row = 0
    
    # Table spanner (full-width header)
    if table_spanner:
        merge_cells_horizontal(table, current_row, 0, num_cols - 1)
        set_cell_text(table.rows[current_row].cells[0], table_spanner, 
                     bold=False, italic=True, align='center', font_size=font_size)
        # Full-width border below table spanner
        set_row_border_bottom(table.rows[current_row])
        current_row += 1
    
    # Column spanners (decked heads - row above column headers)
    if column_spanners:
        spanner_row = table.rows[current_row]
        # First, fill with empty cells
        for col in range(num_cols):
            spanner_row.cells[col].text = ""
        
        # Apply spanners with category-length borders
        for spanner_text, start_col, end_col in column_spanners:
            if start_col < num_cols and end_col < num_cols:
                # Merge cells for multi-column spanners
                if end_col > start_col:
                    merge_cells_horizontal(table, current_row, start_col, end_col)
                # Set spanner text
                set_cell_text(spanner_row.cells[start_col], spanner_text,
                             bold=False, italic=False, align='center', font_size=font_size)
                # Add category-length border under this spanner only
                set_cell_border_bottom(spanner_row.cells[start_col])
        
        current_row += 1
    
    # Column headers row
    header_row = table.rows[current_row]
    for col_idx, header_text in enumerate(column_headers):
        cell = header_row.cells[col_idx]
        # Stub heading (first column) - left aligned
        if col_idx == 0 and stub_heading:
            set_cell_text(cell, stub_heading, bold=False, italic=False, 
                         align='left', font_size=font_size)
        else:
            set_cell_text(cell, header_text, bold=False, italic=False, 
                         align='center', font_size=font_size)
    set_row_border_bottom(header_row)
    current_row += 1
    
    # Data rows
    data_row_start = current_row
    for row_idx, row_data in enumerate(data_rows):
        data_row = table.rows[current_row]
        
        # Add top border for total/summary rows
        if total_row_indices and row_idx in total_row_indices:
            for col_idx in range(num_cols):
                set_cell_border_top(data_row.cells[col_idx])
        
        # Check if this is a cohort subrow (indented with "  Cohort")
        is_cohort_subrow = False
        if row_data and len(row_data) > 0:
            first_col = str(row_data[0])
            is_cohort_subrow = first_col.startswith('  Cohort')
        
        for col_idx, value in enumerate(row_data):
            cell = data_row.cells[col_idx]
            text = str(value) if value is not None and pd.notna(value) else ""
            # First column left-aligned (stub column), others centered to align with headers
            cell.text = text
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT if col_idx == 0 else WD_ALIGN_PARAGRAPH.CENTER
            for run in cell.paragraphs[0].runs:
                run.font.name = 'Times New Roman'
                run.font.size = Pt(font_size)
                # Apply italic to cohort subrows
                if is_cohort_subrow:
                    run.italic = True
        current_row += 1
    
    # Notes
    if note_general:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(6)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run("Note. ")
        run.italic = True
        run.font.name = 'Times New Roman'
        run.font.size = Pt(note_size)
        run = p.add_run(note_general)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(note_size)
    
    if notes_specific:
        for note in notes_specific:
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(0)
            p.paragraph_format.space_after = Pt(0)
            run = p.add_run(note)
            run.font.name = 'Times New Roman'
            run.font.size = Pt(note_size)
    
    if note_probability:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        p.paragraph_format.space_after = Pt(0)
        run = p.add_run(note_probability)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(note_size)
    
    return table_num + 1


def add_apa7_table(doc, table_num, title, df, note_general=None, notes_specific=None, 
                   note_probability=None, compact=False):
    """
    Add APA 7 formatted table (simple version for backward compatibility).
    
    Args:
        compact: If True, reduce spacing for fitting multiple tables per page
    """
    if df is None or df.empty:
        return table_num + 1
    
    # Convert DataFrame to format for advanced function
    column_headers = list(df.columns)
    data_rows = df.values.tolist()
    
    return add_apa7_table_advanced(
        doc, table_num, title, data_rows,
        stub_heading=column_headers[0],
        column_headers=column_headers,
        note_general=note_general,
        notes_specific=notes_specific,
        note_probability=note_probability,
        compact=compact
    )


# =============================================================================
# TABLE GENERATORS
# =============================================================================

def table1_sample_flow(doc, table_num, data_dir, compact=True):
    """Table 1: Sample Flow and Analytic N
    
    With cohort breakdown (2022-23 and 2023-24) shown as grouped subrows.
    """
    sample_file = find_csv(data_dir, "sample_flow.csv")
    
    if sample_file.exists():
        raw = pd.read_csv(sample_file)
        
        # Check if we have cohort breakdown (rows with "  Cohort" prefix)
        has_cohort = any('Cohort' in str(row[0]) for row in raw.values)
        
        if has_cohort:
            # Process with cohort grouping - add visual structure
            data_rows = []
            for _, row in raw.iterrows():
                stage = str(row['Stage'])
                if stage.startswith('  Cohort'):
                    # Cohort subrow - keep indentation, format as italic subgroup
                    data_rows.append([
                        stage,  # Keep "  Cohort 2022-23" with indent
                        fmt_int(row['FASt']),
                        fmt_int(row['Lite_DC']),
                        fmt_int(row['No_Cred']),
                        fmt_int(row['Total'])
                    ])
                else:
                    # Main stage row - bold for group headers
                    data_rows.append([
                        stage,
                        fmt_int(row['FASt']),
                        fmt_int(row['Lite_DC']),
                        fmt_int(row['No_Cred']),
                        fmt_int(row['Total'])
                    ])
        else:
            # No cohort breakdown - simple format
            data_rows = raw.values.tolist()
    else:
        # Placeholder structure
        data_rows = [
            ['BCSSE respondents (baseline)', '—', '—', '—', '—'],
            ['  Cohort 2022-23', '—', '—', '—', '—'],
            ['  Cohort 2023-24', '—', '—', '—', '—'],
            ['Linked to NSSE (follow-up)', '—', '—', '—', '—'],
            ['  Cohort 2022-23', '—', '—', '—', '—'],
            ['  Cohort 2023-24', '—', '—', '—', '—'],
            ['Final analytic sample', '—', '—', '—', '—'],
            ['  Cohort 2022-23', '—', '—', '—', '—'],
            ['  Cohort 2023-24', '—', '—', '—', '—'],
            ['Weighted ESS', '—', '—', '—', '—'],
        ]
    
    # Use column spanners to group treatment columns
    column_spanners = [
        ("Dual Credit Status", 1, 3)  # Spans FASt, Lite_DC, No_Cred
    ]
    
    return add_apa7_table_advanced(
        doc, table_num,
        "Sample Flow and Analytic Sample Size by Treatment Group and Cohort",
        data_rows,
        stub_heading="Selection Stage",
        column_headers=["Selection Stage", "FASt (≥12)", "Lite_DC (1–11)", "No_Cred (0)", "Total"],
        column_spanners=column_spanners,
        note_general=(
            "FASt = ≥12 transferable dual credit units at matriculation; "
            "Lite_DC = 1–11 credits; No_Cred = 0 credits. "
            "Cohort 2022-23 = fall 2022 first-time freshmen; Cohort 2023-24 = fall 2023. "
            "ESS = effective sample size after propensity score overlap weighting."
        ),
        compact=compact
    )


def table2_descriptives(doc, table_num, data_dir, compact=True):
    """Table 2: Descriptive Statistics for All Variables
    
    Uses column spanners to group M/SD and Min/Max (matching Table 5 style).
    Adds Total N row at bottom.
    """
    desc_file = find_csv(data_dir, "descriptives.csv")
    
    total_row_indices = []
    
    if desc_file.exists():
        raw = pd.read_csv(desc_file)
        # Apply readable labels
        raw['Variable'] = raw['Variable'].apply(get_label)
        
        # Build data rows
        data_rows = []
        for _, row in raw.iterrows():
            data_rows.append([
                row['Variable'],
                fmt_int(row['N']),
                fmt(row['M'], 2),
                fmt(row['SD'], 2),
                fmt(row['Min'], 2),
                fmt(row['Max'], 2)
            ])
        
        # Add Total row
        total_n = raw['N'].max()  # Should be same for all
        total_row_indices = [len(data_rows)]
        data_rows.append(['Total N', fmt_int(total_n), '', '', '', ''])
    else:
        # Placeholder structure
        data_rows = [
            ['FASt Status (≥12 credits)', '—', '—', '—', '—', '—'],
            ['Dual Credit Units Earned', '—', '—', '—', '—', '—'],
            ['High School GPA', '—', '—', '—', '—', '—'],
            ['Parent Education Level', '—', '—', '—', '—', '—'],
            ['Pell Grant Recipient', '—', '—', '—', '—', '—'],
            ['AP/IB/College-Level Courses', '—', '—', '—', '—', '—'],
            ['Precalculus or Higher', '—', '—', '—', '—', '—'],
            ['HS Academic Challenge', '—', '—', '—', '—', '—'],
            ['Total N', '—', '', '', '', ''],
        ]
        total_row_indices = [8]
    
    return add_apa7_table_advanced(
        doc, table_num,
        "Descriptive Statistics for Study Variables",
        data_rows,
        stub_heading="Variable",
        column_headers=["Variable", "n", "M", "SD", "Min", "Max"],
        column_spanners=[
            ("Central Tendency", 2, 3),   # Spans M and SD (columns 2-3)
            ("Range", 4, 5),              # Spans Min and Max (columns 4-5)
        ],
        total_row_indices=total_row_indices,
        note_general=(
            "n = sample size; M = mean; SD = standard deviation. "
            "FASt = ≥12 transferable credits; Lite_DC = 1–11 credits; No_Cred = 0 credits."
        ),
        compact=compact
    )


def table3_missing_data(doc, table_num, data_dir, compact=True):
    """Table 3: Missing Data Patterns and Handling"""
    
    missing_file = find_csv(data_dir, "missing_data.csv")
    if missing_file.exists():
        raw = pd.read_csv(missing_file)
        # Check if it's the simple 2-column format or full format
        if 'Missing_Pct' in raw.columns:
            # Simple format from R - apply labels and format
            data_rows = []
            for _, row in raw.iterrows():
                data_rows.append([
                    get_label(row['Variable']),
                    fmt_pct(row['Missing_Pct'], 1)
                ])
        else:
            data_rows = raw.values.tolist()
    else:
        data_rows = [
            ['FASt Status (≥12 credits)', '0.0%'],
            ['Dual Credit Units Earned', '0.0%'],
            ['Emotional Distress items', '—'],
            ['Quality of Engagement items', '—'],
            ['Developmental Adjustment items', '—'],
            ['High School GPA', '—'],
            ['Pell status', '—'],
        ]
    
    return add_apa7_table_advanced(
        doc, table_num,
        "Missing Data Rates by Variable",
        data_rows,
        stub_heading="Variable",
        column_headers=["Variable", "Missing %"],
        note_general=(
            "Missing data assumed MAR (missing at random) and handled via FIML "
            "(full information maximum likelihood) in lavaan. "
            "Percentages reflect proportion of cases with missing values on each variable."
        ),
        compact=compact
    )


def table4_ps_model(doc, table_num, data_dir, compact=True):
    """Table 4: Propensity Score Model Specification
    
    Uses column spanners for B/SE grouping (matching Table 5 style).
    """
    ps_file = find_csv(data_dir, "ps_model.csv")
    
    if ps_file.exists():
        df = pd.read_csv(ps_file)
        # Apply readable labels to Covariate column
        if 'Covariate' in df.columns:
            df['Covariate'] = df['Covariate'].apply(get_label)
        
        # Build data rows
        data_rows = []
        for _, row in df.iterrows():
            data_rows.append([
                row['Covariate'],
                fmt(row.get('B', None), 3),
                fmt(row.get('SE', None), 3),
                fmt(row.get('OR', None), 3)
            ])
    else:
        data_rows = [
            ['Cohort (2022–23 vs 2023–24)', '—', '—', '—'],
            ['High School GPA (centered)', '—', '—', '—'],
            ['Parental Education (centered)', '—', '—', '—'],
            ['Pell Grant Recipient', '—', '—', '—'],
            ['AP/IB/College-Level Course', '—', '—', '—'],
            ['Precalculus Completion', '—', '—', '—'],
            ['High School Academic Challenge (centered)', '—', '—', '—'],
        ]
    
    return add_apa7_table_advanced(
        doc, table_num,
        "Propensity Score Model: Logistic Regression Predicting FASt Status",
        data_rows,
        stub_heading="Covariate",
        column_headers=["Covariate", "B", "SE", "OR"],
        column_spanners=[
            ("Estimate", 1, 2),  # Spans B and SE (columns 1-2)
        ],
        note_general=(
            "Treatment: FASt = 1 (≥12 credits), non-FASt = 0 (Lite_DC + No_Cred). "
            "B = unstandardized coefficient; SE = standard error; OR = odds ratio. "
            "All continuous predictors mean-centered. Model used for computing overlap (ATO) weights."
        ),
        compact=compact
    )


def table5_balance(doc, table_num, data_dir, compact=True):
    """Table 5: Covariate Balance Before and After Weighting
    
    Uses APA 7 column spanners (decked heads) for Pre/Post groupings.
    Total rows (Mean |SMD|, Max |SMD|) get horizontal separator.
    """
    balance_file = find_csv(data_dir, "balance.csv")
    
    total_row_indices = []
    
    if balance_file.exists():
        raw = pd.read_csv(balance_file)
        # Build data rows with readable labels
        data_rows = []
        for idx, row in raw.iterrows():
            covariate_raw = row['Covariate']
            covariate = get_label(covariate_raw)
            
            # Mark summary rows for separator
            if covariate_raw in ['Mean |SMD|', 'Max |SMD|'] or 'Mean' in str(covariate_raw) or 'Max' in str(covariate_raw):
                if len(data_rows) > 0 and covariate_raw == 'Mean |SMD|':
                    total_row_indices.append(len(data_rows))
            
            data_rows.append([
                covariate,
                fmt(row['SMD_Pre'], 3),
                fmt(row['VR_Pre'], 3),
                fmt(row['SMD_Post'], 3),
                fmt(row['VR_Post'], 3)
            ])
    else:
        data_rows = [
            ['Cohort', '—', '—', '—', '—'],
            ['High School GPA', '—', '—', '—', '—'],
            ['Parental Education', '—', '—', '—', '—'],
            ['Pell Grant Recipient', '—', '—', '—', '—'],
            ['AP/IB/College-Level Course', '—', '—', '—', '—'],
            ['Precalculus Completion', '—', '—', '—', '—'],
            ['High School Academic Challenge', '—', '—', '—', '—'],
            ['Mean |SMD|', '—', '—', '—', '—'],
            ['Max |SMD|', '—', '—', '—', '—'],
        ]
        total_row_indices = [7]  # Mean |SMD| row gets separator
    
    return add_apa7_table_advanced(
        doc, table_num,
        "Covariate Balance Before and After Propensity Score Weighting",
        data_rows,
        stub_heading="Covariate",
        column_headers=["Covariate", "SMD", "VR", "SMD", "VR"],
        column_spanners=[
            ("Pre-weighting", 1, 2),   # Spans columns 1-2
            ("Post-weighting", 3, 4),  # Spans columns 3-4
        ],
        total_row_indices=total_row_indices,
        note_general=(
            "SMD = standardized mean difference (FASt − non-FASt). "
            "VR = variance ratio (FASt / non-FASt). Non-FASt = Lite_DC + No_Cred combined. "
            "Balance achieved when |SMD| < 0.10 and VR between 0.80–1.25."
        ),
        compact=compact
    )


def table6_weights(doc, table_num, data_dir, compact=True):
    """Table 6: Weight Diagnostics
    
    Total row gets horizontal separator.
    Groups weight distribution columns under spanner.
    """
    weights_file = find_csv(data_dir, "weight_diagnostics.csv")
    
    total_row_indices = []
    
    if weights_file.exists():
        df = pd.read_csv(weights_file)
        
        # Standardize group names with credit indicators
        group_mapping = {
            'Overall': 'Total',
            'FASt (treated)': 'FASt (≥12)',
            'Non-FASt (control)': 'Non-FASt (<12)',
        }
        df['Group'] = df['Group'].replace(group_mapping)
        
        # Format numeric columns
        df['N'] = df['N'].apply(fmt_int)
        for col in ['Min', 'P5', 'Median', 'P95', 'Max']:
            df[col] = df[col].apply(lambda x: fmt(x, 3))
        df['ESS'] = df['ESS'].apply(lambda x: fmt(x, 1))
        
        # Find Total row for separator
        for idx, row in df.iterrows():
            if 'Total' in str(row.get('Group', '')):
                total_row_indices.append(idx)
                break
        
        data_rows = df.values.tolist()
    else:
        data_rows = [
            ['FASt (≥12)', '—', '—', '—', '—', '—', '—', '—'],
            ['Non-FASt (<12)', '—', '—', '—', '—', '—', '—', '—'],
            ['Total', '—', '—', '—', '—', '—', '—', '—'],
        ]
        total_row_indices = [2]  # Total row
    
    return add_apa7_table_advanced(
        doc, table_num,
        "Propensity Score Overlap Weight Diagnostics",
        data_rows,
        stub_heading="Group",
        column_headers=["Group", "n", "Min", "P5", "Mdn", "P95", "Max", "ESS"],
        column_spanners=[
            ("Weight Distribution", 2, 6),  # Spans Min through Max (columns 2-6)
        ],
        total_row_indices=total_row_indices,
        note_general=(
            "FASt = ≥12 credits; Non-FASt = Lite_DC (1–11) + No_Cred (0). "
            "Overlap weights: w = 1 − PS for treated, w = PS for control. "
            "ESS = effective sample size = (Σw)²/Σw². "
            "P5/P95 = 5th/95th percentiles; Mdn = median."
        ),
        compact=compact
    )


def table7_cfa(doc, table_num, data_dir, compact=True):
    """Table 7: Measurement Model (CFA) Results
    
    Uses column spanners for Loading (λ/SE) grouping (matching Table 5 style).
    """
    cfa_file = find_csv(data_dir, "cfa_results.csv")
    if cfa_file.exists():
        df = pd.read_csv(cfa_file)
        # Apply labels to Item/Factor column if present
        if 'Item/Factor' in df.columns:
            def format_item(x):
                if not x:
                    return x
                if x.startswith(' '):
                    return '  ' + get_label(x.strip())
                return get_label(x.strip())
            df['Item/Factor'] = df['Item/Factor'].apply(format_item)
        data_rows = df.values.tolist()
    else:
        data_rows = [
            ['Emotional Distress (EmoDiss)', '', '', '', ''],
            ['  Difficulty: Academic Work', '—', '—', '', ''],
            ['  Difficulty: Loneliness', '—', '—', '', ''],
            ['  Difficulty: Mental Health', '—', '—', '', ''],
            ['  Difficulty: Exhaustion', '—', '—', '', ''],
            ['  Difficulty: Sleep Problems', '—', '—', '', ''],
            ['  Difficulty: Financial Concerns', '—', '—', '', ''],
            ['  Factor ω', '', '', '—', ''],
            ['Quality of Engagement (QualEngag)', '', '', '', ''],
            ['  Quality: Student Interactions', '—', '—', '', ''],
            ['  Quality: Advisor Interactions', '—', '—', '', ''],
            ['  Quality: Faculty Interactions', '—', '—', '', ''],
            ['  Quality: Staff Interactions', '—', '—', '', ''],
            ['  Quality: Admin Interactions', '—', '—', '', ''],
            ['  Factor ω', '', '', '—', ''],
        ]
    
    return add_apa7_table_advanced(
        doc, table_num,
        "Confirmatory Factor Analysis: Standardized Factor Loadings",
        data_rows,
        stub_heading="Item/Factor",
        column_headers=["Item/Factor", "λ", "SE", "ω", "AVE"],
        column_spanners=[
            ("Loading", 1, 2),      # Spans λ and SE (columns 1-2)
            ("Reliability", 3, 4),  # Spans ω and AVE (columns 3-4)
        ],
        note_general=(
            "λ = standardized factor loading; SE = standard error. "
            "ω = McDonald's omega reliability; AVE = average variance extracted. "
            "All loadings p < .001. Marker-variable identification used."
        ),
        compact=compact
    )


def table8_invariance(doc, table_num, data_dir, compact=True):
    """Table 8: Measurement Invariance Tests
    
    Uses APA 7 column spanners for Fit Indices and Model Comparison.
    """
    inv_file = find_csv(data_dir, "invariance.csv")
    
    if inv_file.exists():
        raw = pd.read_csv(inv_file)
        data_rows = raw.values.tolist()
    else:
        data_rows = [
            ['Configural', '—', '—', '—', '—', '—', '—', '—'],
            ['Metric', '—', '—', '—', '—', '—', '—', '—'],
            ['Scalar', '—', '—', '—', '—', '—', '—', '—'],
        ]
    
    return add_apa7_table_advanced(
        doc, table_num,
        "Measurement Invariance Tests Across Treatment Groups",
        data_rows,
        stub_heading="Model",
        column_headers=["Model", "χ²", "df", "CFI", "RMSEA", "ΔCFI", "ΔRMSEA", "Decision"],
        column_spanners=[
            ("Absolute Fit", 1, 4),
            ("Δ Comparison", 5, 6),
        ],
        note_general=(
            "Invariance tested for EmoDiss, QualEngag, and DevAdj factors. "
            "ΔCFI > −.010 and ΔRMSEA < .015 indicate acceptable invariance. "
            "Scalar invariance required for latent mean comparisons."
        ),
        compact=compact
    )


def table9_model_fit(doc, table_num, data_dir, B, ci_type, compact=True):
    """Table 9: Structural Model Fit
    
    Uses column spanners for fit index groupings (matching Table 5 style).
    """
    # Try multiple file formats
    fit_file = find_csv(data_dir, "model_fit.csv")
    pe_fit_file = find_csv(data_dir, "point_estimates_fit_indices.csv")
    fit_txt_file = data_dir / "structural_fitMeasures.txt"
    
    data_rows = None
    
    if fit_file.exists():
        raw = pd.read_csv(fit_file)
        data_rows = raw.values.tolist()
    elif pe_fit_file.exists():
        # Read fit indices from R output and format
        raw = pd.read_csv(pe_fit_file)
        # Format the RMSEA CI
        rmsea_ci = f"{fmt(raw.get('rmsea', [None])[0], 3)} [{fmt(raw.get('rmsea.ci.lower', [None])[0], 3)}, {fmt(raw.get('rmsea.ci.upper', [None])[0], 3)}]"
        data_rows = [
            ['Full structural model',
             '—',  # chi-square not always reported with MLR
             '—',  # df
             fmt(raw.get('cfi', [None])[0], 3),
             fmt(raw.get('tli', [None])[0], 3),
             rmsea_ci,
             fmt(raw.get('srmr', [None])[0], 3)],
        ]
    elif fit_txt_file.exists():
        # Read the structural_fitMeasures.txt format (two columns: measure, value)
        try:
            fit_df = pd.read_csv(fit_txt_file, sep='\t')
            # Convert to dict for easy access
            fit_dict = dict(zip(fit_df['measure'], fit_df['value']))
            
            # Get values (prefer robust/scaled if available)
            cfi = fit_dict.get('cfi.robust', fit_dict.get('cfi.scaled', fit_dict.get('cfi', None)))
            tli = fit_dict.get('tli.robust', fit_dict.get('tli.scaled', fit_dict.get('tli', None)))
            rmsea = fit_dict.get('rmsea.robust', fit_dict.get('rmsea.scaled', fit_dict.get('rmsea', None)))
            srmr = fit_dict.get('srmr', None)
            chisq = fit_dict.get('chisq', None)
            df = fit_dict.get('df', None)
            
            # Format RMSEA (no CI available in this format)
            rmsea_str = fmt(rmsea, 3)
            
            data_rows = [
                ['Full structural model',
                 fmt(chisq, 2),
                 fmt_int(df) if df else '—',
                 fmt(cfi, 3),
                 fmt(tli, 3),
                 rmsea_str,
                 fmt(srmr, 3)],
            ]
        except Exception as e:
            print(f"Warning: Could not parse {fit_txt_file}: {e}")
    
    if data_rows is None:
        data_rows = [
            ['Measurement model (CFA)', '—', '—', '—', '—', '—', '—'],
            ['Structural model (full)', '—', '—', '—', '—', '—', '—'],
            ['Alternative: No moderation', '—', '—', '—', '—', '—', '—'],
        ]
    
    return add_apa7_table_advanced(
        doc, table_num,
        "Structural Equation Model Fit Indices",
        data_rows,
        stub_heading="Model",
        column_headers=["Model", "χ²", "df", "CFI", "TLI", "RMSEA [90% CI]", "SRMR"],
        column_spanners=[
            ("Chi-Square", 1, 2),        # Spans χ² and df
            ("Incremental Fit", 3, 4),   # Spans CFI and TLI
            ("Absolute Fit", 5, 6),      # Spans RMSEA and SRMR
        ],
        note_general=(
            f"Estimator: MLR with robust standard errors. Missing data: FIML. "
            f"Bootstrap: B = {B:,} replicates. "
            "Acceptable fit: CFI/TLI ≥ .90, RMSEA ≤ .08, SRMR ≤ .08."
        ),
        compact=compact
    )


def table10_structural_paths(doc, table_num, bootstrap_df, B, ci_type, data_dir, compact=False):
    """Table 10: Structural Path Coefficients
    
    Groups B/SE under Unstandardized spanner, β under Standardized.
    """
    ci_label = {'bca': 'BCa', 'perc': 'percentile', 'norm': 'normal'}.get(ci_type.lower(), ci_type)
    
    struct_params = ['a1', 'a1z', 'a2', 'a2z', 'b1', 'b2', 'c', 'cz']
    paths = bootstrap_df[bootstrap_df['parameter'].isin(struct_params)].copy()
    
    if paths.empty:
        return table_num
    
    # Load standardized coefficients
    std_map = load_standardized_coefficients(data_dir)
    
    labels = {
        'a1': 'a₁: X → M₁ (FASt → EmoDiss)',
        'a1z': 'a₁z: X×Z → M₁ (interaction)',
        'a2': 'a₂: X → M₂ (FASt → QualEngag)',
        'a2z': 'a₂z: X×Z → M₂ (interaction)',
        'b1': 'b₁: M₁ → Y (EmoDiss → DevAdj)',
        'b2': 'b₂: M₂ → Y (QualEngag → DevAdj)',
        'c': "c′: X → Y (direct effect)",
        'cz': "c′z: X×Z → Y (direct × dose)"
    }
    
    paths['Path'] = paths['parameter'].map(labels)
    paths['B'] = paths['est'].apply(lambda x: fmt(x, 3))
    paths['SE'] = paths['boot_se'].apply(lambda x: fmt(x, 3))
    paths['β'] = paths['parameter'].apply(lambda x: fmt(std_map.get(x, None), 3))
    paths['95% CI'] = paths.apply(lambda r: fmt_ci(r['ci_lower'], r['ci_upper'], 3), axis=1)
    paths['Sig'] = paths['sig'].apply(lambda x: '*' if x else '')
    
    data_rows = paths[['Path', 'B', 'SE', 'β', '95% CI', 'Sig']].values.tolist()
    
    return add_apa7_table_advanced(
        doc, table_num,
        "Structural Path Coefficients From Bootstrap Analysis",
        data_rows,
        stub_heading="Path",
        column_headers=["Path", "B", "SE", "β", "95% CI", "Sig"],
        column_spanners=[
            ("Unstandardized", 1, 2),  # Spans B and SE
            ("Std.", 3, 3),            # β alone
        ],
        note_general=(
            "X = FASt status; Z = credit dose (centered); M₁ = Emotional Distress; "
            "M₂ = Quality of Engagement; Y = Developmental Adjustment. "
            "B = unstandardized coefficient; β = fully standardized coefficient."
        ),
        note_probability=f"*95% {ci_label} CI excludes zero. Bootstrap B = {B:,}.",
        compact=compact
    )


def table11_indirect_effects(doc, table_num, bootstrap_df, B, ci_type, data_dir, compact=False):
    """Table 11: Indirect Effects
    
    Groups B/SE under Unstandardized, β under Standardized.
    """
    ci_label = {'bca': 'BCa', 'perc': 'percentile', 'norm': 'normal'}.get(ci_type.lower(), ci_type)
    
    # Load standardized coefficients
    std_map = load_standardized_coefficients(data_dir)
    
    # Get indirect effects at mean (z_mid)
    ind_params = ['ind_EmoDiss_z_mid', 'ind_QualEngag_z_mid']
    ind = bootstrap_df[bootstrap_df['parameter'].isin(ind_params)].copy()
    
    # Also get total indirect and total effect if available
    total_params = ['total_z_mid']
    totals = bootstrap_df[bootstrap_df['parameter'].isin(total_params)].copy()
    
    # Direct effect
    direct = bootstrap_df[bootstrap_df['parameter'] == 'c'].copy()
    
    rows = []
    
    # Direct effect
    if not direct.empty:
        r = direct.iloc[0]
        rows.append(['Direct effect (c′)', fmt(r['est'], 3), fmt(r['boot_se'], 3), 
                     fmt(std_map.get('c', None), 3),
                     fmt_ci(r['ci_lower'], r['ci_upper'], 3), '*' if r['sig'] else ''])
    
    # Indirect via EmoDiss
    ind_emo = bootstrap_df[bootstrap_df['parameter'] == 'ind_EmoDiss_z_mid']
    if not ind_emo.empty:
        r = ind_emo.iloc[0]
        rows.append(['Indirect via EmoDiss (a₁×b₁)', fmt(r['est'], 3), fmt(r['boot_se'], 3),
                     fmt(std_map.get('ind_EmoDiss_z_mid', None), 3),
                     fmt_ci(r['ci_lower'], r['ci_upper'], 3), '*' if r['sig'] else ''])
    
    # Indirect via QualEngag
    ind_qual = bootstrap_df[bootstrap_df['parameter'] == 'ind_QualEngag_z_mid']
    if not ind_qual.empty:
        r = ind_qual.iloc[0]
        rows.append(['Indirect via QualEngag (a₂×b₂)', fmt(r['est'], 3), fmt(r['boot_se'], 3),
                     fmt(std_map.get('ind_QualEngag_z_mid', None), 3),
                     fmt_ci(r['ci_lower'], r['ci_upper'], 3), '*' if r['sig'] else ''])
    
    # Total indirect (sum of both)
    total_ind_idx = None
    if not ind_emo.empty and not ind_qual.empty:
        total_ind = ind_emo.iloc[0]['est'] + ind_qual.iloc[0]['est']
        total_ind_std = (std_map.get('ind_EmoDiss_z_mid', 0) or 0) + (std_map.get('ind_QualEngag_z_mid', 0) or 0)
        total_ind_idx = len(rows)
        rows.append(['Total indirect', fmt(total_ind, 3), '—', fmt(total_ind_std, 3) if total_ind_std else '—', '—', ''])
    
    # Total effect
    if not totals.empty:
        r = totals.iloc[0]
        rows.append(['Total effect', fmt(r['est'], 3), fmt(r['boot_se'], 3),
                     fmt(std_map.get('total_z_mid', None), 3),
                     fmt_ci(r['ci_lower'], r['ci_upper'], 3), '*' if r['sig'] else ''])
    
    if not rows:
        return table_num
    
    # Mark total rows for separator
    total_row_indices = [total_ind_idx] if total_ind_idx is not None else []
    
    return add_apa7_table_advanced(
        doc, table_num,
        "Direct, Indirect, and Total Effects at Mean Credit Dose",
        rows,
        stub_heading="Effect",
        column_headers=["Effect", "B", "SE", "β", "95% CI", "Sig"],
        column_spanners=[
            ("Unstandardized", 1, 2),  # Spans B and SE
            ("Std.", 3, 3),            # β alone
        ],
        total_row_indices=total_row_indices,
        note_general=(
            "Effects evaluated at mean credit dose (Z = 0). "
            "Indirect effect = product of a-path and b-path coefficients. "
            "B = unstandardized; β = fully standardized."
        ),
        note_probability=f"*95% {ci_label} bootstrap CI excludes zero. B = {B:,}.",
        compact=compact
    )


def table12_conditional_indirect(doc, table_num, bootstrap_df, B, ci_type, data_dir, compact=False):
    """Table 12: Conditional Indirect Effects and Index of Moderated Mediation
    
    Groups B/SE under Unstandardized, β under Standardized.
    """
    ci_label = {'bca': 'BCa', 'perc': 'percentile', 'norm': 'normal'}.get(ci_type.lower(), ci_type)
    
    # Load standardized coefficients
    std_map = load_standardized_coefficients(data_dir)
    
    # Build data rows with table spanner sections
    rows = []
    
    # Section: Via Emotional Distress
    rows.append(['Via Emotional Distress (a₁ × b₁)', '', '', '', '', ''])
    for level, label in [('low', '  At −1 SD credit dose'), ('mid', '  At mean credit dose'), ('high', '  At +1 SD credit dose')]:
        param = f'ind_EmoDiss_z_{level}'
        r = bootstrap_df[bootstrap_df['parameter'] == param]
        if not r.empty:
            r = r.iloc[0]
            rows.append([label, fmt(r['est'], 3), fmt(r['boot_se'], 3),
                         fmt(std_map.get(param, None), 3),
                         fmt_ci(r['ci_lower'], r['ci_upper'], 3), '*' if r['sig'] else ''])
    
    # IMM for EmoDiss
    imm_emo = bootstrap_df[bootstrap_df['parameter'] == 'index_MM_EmoDiss']
    if not imm_emo.empty:
        r = imm_emo.iloc[0]
        rows.append(['  Index of moderated mediation', fmt(r['est'], 3), fmt(r['boot_se'], 3),
                     fmt(std_map.get('index_MM_EmoDiss', None), 3),
                     fmt_ci(r['ci_lower'], r['ci_upper'], 3), '*' if r['sig'] else ''])
    
    # Section: Via Quality of Engagement
    rows.append(['Via Quality of Engagement (a₂ × b₂)', '', '', '', '', ''])
    for level, label in [('low', '  At −1 SD credit dose'), ('mid', '  At mean credit dose'), ('high', '  At +1 SD credit dose')]:
        param = f'ind_QualEngag_z_{level}'
        r = bootstrap_df[bootstrap_df['parameter'] == param]
        if not r.empty:
            r = r.iloc[0]
            rows.append([label, fmt(r['est'], 3), fmt(r['boot_se'], 3),
                         fmt(std_map.get(param, None), 3),
                         fmt_ci(r['ci_lower'], r['ci_upper'], 3), '*' if r['sig'] else ''])
    
    # IMM for QualEngag
    imm_qual = bootstrap_df[bootstrap_df['parameter'] == 'index_MM_QualEngag']
    if not imm_qual.empty:
        r = imm_qual.iloc[0]
        rows.append(['  Index of moderated mediation', fmt(r['est'], 3), fmt(r['boot_se'], 3),
                     fmt(std_map.get('index_MM_QualEngag', None), 3),
                     fmt_ci(r['ci_lower'], r['ci_upper'], 3), '*' if r['sig'] else ''])
    
    if len(rows) <= 2:  # Only section headers
        return table_num
    
    return add_apa7_table_advanced(
        doc, table_num,
        "Conditional Indirect Effects and Index of Moderated Mediation",
        rows,
        stub_heading="Effect",
        column_headers=["Effect", "B", "SE", "β", "95% CI", "Sig"],
        column_spanners=[
            ("Unstandardized", 1, 2),  # Spans B and SE
            ("Std.", 3, 3),            # β alone
        ],
        note_general=(
            "Conditional indirect effects evaluated at −1 SD, mean, and +1 SD of credit dose. "
            "IMM = index of moderated mediation = a₁z × b (slope of indirect effect on moderator). "
            "B = unstandardized; β = fully standardized."
        ),
        note_probability=f"*95% {ci_label} bootstrap CI excludes zero. B = {B:,} replicates.",
        compact=compact
    )


def table13_robustness(doc, table_num, data_dir, compact=True):
    """Table 13: Robustness Checks"""
    
    robust_file = find_csv(data_dir, "robustness.csv")
    if robust_file.exists():
        df = pd.read_csv(robust_file)
    else:
        df = pd.DataFrame([
            ['a₁ (X → M₁)', '—', '—', '—', '—'],
            ['a₂ (X → M₂)', '—', '—', '—', '—'],
            ['b₁ (M₁ → Y)', '—', '—', '—', '—'],
            ['b₂ (M₂ → Y)', '—', '—', '—', '—'],
            ['c′ (direct)', '—', '—', '—', '—'],
            ['IMM (EmoDiss)', '—', '—', '—', '—'],
            ['IMM (QualEngag)', '—', '—', '—', '—'],
        ], columns=['Parameter', 'Weighted B', 'Unweighted B', 'IPTW B', 'Difference'])
    
    return add_apa7_table(
        doc, table_num,
        "Robustness Checks: Comparison Across Estimation Methods",
        df,
        note_general=(
            "Weighted = overlap (ATO) weights; Unweighted = no propensity adjustment; "
            "IPTW = inverse probability of treatment weights. "
            "Difference = |Weighted − Unweighted|."
        ),
        compact=compact
    )


def table14_serial_mediation(doc, table_num, data_dir, compact=True):
    """Table 14: Serial Mediation Results (Exploratory)"""
    
    serial_file = find_csv(data_dir, "serial_mediation.csv")
    if serial_file.exists():
        df = pd.read_csv(serial_file)
        # Convert Significant column to * notation
        if 'Significant' in df.columns:
            df['Sig'] = df['Significant'].apply(lambda x: '*' if str(x).lower() in ['yes', 'true', '1'] else '')
            df['Estimate'] = df['Estimate'].apply(lambda x: fmt(x, 4))
            df['SE'] = df['SE'].apply(lambda x: fmt(x, 4))
            df['95% CI'] = df.apply(lambda r: fmt_ci(r['CI_Lower'], r['CI_Upper'], 4), axis=1)
            df = df[['Parameter', 'Estimate', 'SE', '95% CI', 'Sig']]
    else:
        df = pd.DataFrame([
            ['d (EmoDiss → QualEngag)', '—', '—', '—', ''],
            ['Serial Indirect (Low Dose)', '—', '—', '—', ''],
            ['Serial Indirect (Mean Dose)', '—', '—', '—', ''],
            ['Serial Indirect (High Dose)', '—', '—', '—', ''],
            ['Index of Moderated Mediation (Serial)', '—', '—', '—', ''],
        ], columns=['Parameter', 'Estimate', 'SE', '95% CI', 'Sig'])
    
    return add_apa7_table(
        doc, table_num,
        "Serial Mediation Results (Exploratory Analysis)",
        df,
        note_general=(
            "Serial mediation path: FASt → EmoDiss → QualEngag → DevAdj. "
            "The d parameter represents the direct effect of EmoDiss on QualEngag. "
            "Serial indirect effects evaluated at low (-1 SD), mean, and high (+1 SD) credit dose. "
            "Index of moderated mediation (IMM) indicates dose-dependent change in serial indirect effect."
        ),
        notes_specific=[
            "Exploratory analysis examining sequential mediation. "
            "Significant = 95% CI excludes zero."
        ],
        compact=compact
    )


# =============================================================================
# MAIN
# =============================================================================

def main():
    parser = argparse.ArgumentParser(description='Build complete dissertation tables')
    parser.add_argument('--outdir', type=str, required=True, help='Input directory with results data')
    parser.add_argument('--out', type=str, default=None, help='Output directory for docx (default: same as outdir)')
    parser.add_argument('--B', type=int, default=2000, help='Bootstrap replicates')
    parser.add_argument('--ci_type', type=str, default='perc', help='CI type: bca, perc, norm')
    args = parser.parse_args()
    
    outdir = Path(args.outdir)
    out_docx_dir = Path(args.out) if args.out else outdir
    out_docx_dir.mkdir(parents=True, exist_ok=True)
    
    # Find bootstrap results
    boot_csv = outdir / "bootstrap_results.csv"
    if not boot_csv.exists():
        boot_csv = outdir / "pooled" / "bootstrap_results.csv"
    
    bootstrap_df = None
    if boot_csv.exists():
        bootstrap_df = pd.read_csv(boot_csv)
        if 'se' in bootstrap_df.columns and 'boot_se' not in bootstrap_df.columns:
            bootstrap_df['boot_se'] = bootstrap_df['se']
        print(f"Loaded bootstrap results: {boot_csv}")
    else:
        print(f"Warning: No bootstrap_results.csv found in {outdir}")
        bootstrap_df = pd.DataFrame()
    
    # Create document
    doc = Document()
    style = doc.styles['Normal']
    style.font.name = 'Times New Roman'
    style.font.size = Pt(12)
    
    # Set narrow margins for more content per page
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)
    
    table_num = 1
    
    # =========================================================================
    # PART 1: SAMPLE & DESCRIPTIVES (Tables 1-3) - Compact, fit on 1-2 pages
    # =========================================================================
    p = doc.add_paragraph()
    p.add_run("PART I: SAMPLE AND DATA QUALITY").bold = True
    
    table_num = table1_sample_flow(doc, table_num, outdir, compact=True)
    table_num = table2_descriptives(doc, table_num, outdir, compact=True)
    table_num = table3_missing_data(doc, table_num, outdir, compact=True)
    
    doc.add_page_break()
    
    # =========================================================================
    # PART 2: PROPENSITY SCORE WEIGHTING (Tables 4-6)
    # =========================================================================
    p = doc.add_paragraph()
    p.add_run("PART II: PROPENSITY SCORE WEIGHTING DIAGNOSTICS").bold = True
    
    table_num = table4_ps_model(doc, table_num, outdir, compact=True)
    table_num = table5_balance(doc, table_num, outdir, compact=True)
    table_num = table6_weights(doc, table_num, outdir, compact=True)
    
    doc.add_page_break()
    
    # =========================================================================
    # PART 3: MEASUREMENT MODEL (Tables 7-9)
    # =========================================================================
    p = doc.add_paragraph()
    p.add_run("PART III: MEASUREMENT MODEL").bold = True
    
    table_num = table7_cfa(doc, table_num, outdir, compact=True)
    table_num = table8_invariance(doc, table_num, outdir, compact=True)
    table_num = table9_model_fit(doc, table_num, outdir, args.B, args.ci_type, compact=True)
    
    doc.add_page_break()
    
    # =========================================================================
    # PART 4: STRUCTURAL RESULTS (Tables 10-12) - Full size, key results
    # =========================================================================
    p = doc.add_paragraph()
    p.add_run("PART IV: STRUCTURAL MODEL RESULTS").bold = True
    
    if not bootstrap_df.empty:
        table_num = table10_structural_paths(doc, table_num, bootstrap_df, args.B, args.ci_type, outdir, compact=False)
        table_num = table11_indirect_effects(doc, table_num, bootstrap_df, args.B, args.ci_type, outdir, compact=False)
        table_num = table12_conditional_indirect(doc, table_num, bootstrap_df, args.B, args.ci_type, outdir, compact=False)
    else:
        p = doc.add_paragraph("Bootstrap results not found. Tables 10-12 require bootstrap_results.csv.")
    
    doc.add_page_break()
    
    # =========================================================================
    # PART 5: ROBUSTNESS (Table 13) + SERIAL MEDIATION (Table 14)
    # =========================================================================
    p = doc.add_paragraph()
    p.add_run("PART V: ROBUSTNESS AND EXPLORATORY ANALYSES").bold = True
    
    table_num = table13_robustness(doc, table_num, outdir, compact=True)
    
    # Table 14: Serial mediation (if available)
    serial_dir = outdir / "serial_mediation"
    if serial_dir.exists():
        table_num = table14_serial_mediation(doc, table_num, serial_dir, compact=True)
    
    # Save to specified output directory
    out_docx = out_docx_dir / "Dissertation_Tables.docx"
    doc.save(str(out_docx))
    
    print(f"\n{'='*60}")
    print(f"Wrote: {out_docx}")
    print(f"{'='*60}")
    print(f"\nTables generated (APA 7 format):")
    print(f"  PART I:   Tables 1-3  (Sample, Descriptives, Missing)")
    print(f"  PART II:  Tables 4-6  (PS Model, Balance, Weights)")
    print(f"  PART III: Tables 7-9  (CFA, Invariance, Fit)")
    print(f"  PART IV:  Tables 10-12 (Paths, Indirect, Conditional)")
    print(f"  PART V:   Table 13    (Robustness)")
    if serial_dir.exists():
        print(f"            Table 14    (Serial Mediation - Exploratory)")
    print(f"\nMetadata: B = {args.B:,} | CI = {args.ci_type}")
    print(f"\nNote: Tables with '—' placeholders need data files in {outdir}/")
    print(f"      Required: sample_flow.csv, descriptives.csv, missing_data.csv,")
    print(f"                ps_model.csv, balance.csv, weight_diagnostics.csv,")
    print(f"                cfa_results.csv, invariance.csv, model_fit.csv, robustness.csv")
    if serial_dir.exists():
        print(f"      Optional: serial_mediation/serial_mediation.csv (for Table 14)")


if __name__ == "__main__":
    main()
