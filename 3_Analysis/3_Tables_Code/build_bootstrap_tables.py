#!/usr/bin/env python3
"""
Build APA 7 Bootstrap Tables (called by R pipeline)

Generates Tables 10-12 for structural results:
- Table 10: Structural path coefficients
- Table 11: Direct, indirect, and total effects
- Table 12: Conditional indirect effects and IMM

APA 7 Table Formatting:
- Table number: Bold
- Title: Italic, title case
- Stub column (first): Left-aligned (row labels)
- Data columns: Right-aligned (numeric values)
- Borders: Top, below header, bottom only (no vertical lines)
- Notes: "Note." in italics, then regular text

Usage:
    python scripts/build_bootstrap_tables.py --csv <path> --B <int> --ci_type <str>
"""

import argparse
from pathlib import Path
import pandas as pd
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml


def fmt(x, nd=3):
    try:
        return f"{float(x):.{nd}f}" if pd.notna(x) else "—"
    except:
        return str(x) if pd.notna(x) else "—"


def fmt_ci(lo, hi, nd=3):
    return f"[{fmt(lo, nd)}, {fmt(hi, nd)}]"


def set_apa_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr or parse_xml(r'<w:tblPr %s/>' % nsdecls('w'))
    borders = parse_xml(
        r'<w:tblBorders %s>'
        r'<w:top w:val="single" w:sz="4" w:color="000000"/>'
        r'<w:bottom w:val="single" w:sz="4" w:color="000000"/>'
        r'<w:insideH w:val="nil"/><w:insideV w:val="nil"/>'
        r'<w:left w:val="nil"/><w:right w:val="nil"/>'
        r'</w:tblBorders>' % nsdecls('w'))
    tblPr.append(borders)


def header_border(row):
    for cell in row.cells:
        tcPr = cell._tc.get_or_add_tcPr()
        tcPr.append(parse_xml(r'<w:tcBorders %s><w:bottom w:val="single" w:sz="4" w:color="000000"/></w:tcBorders>' % nsdecls('w')))


def set_cell_border_bottom(cell):
    """Add bottom border to a single cell (for category-length spanner borders)."""
    tcPr = cell._tc.get_or_add_tcPr()
    tcPr.append(parse_xml(r'<w:tcBorders %s><w:bottom w:val="single" w:sz="4" w:color="000000"/></w:tcBorders>' % nsdecls('w')))


def add_table(doc, num, title, df, note, note_p=None, column_spanners=None):
    # Number (bold)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(f"Table {num}")
    run.bold = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    
    # Title (italic)
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(title)
    run.italic = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(12)
    
    if df is not None and not df.empty:
        # Add extra row for column spanners if provided
        num_header_rows = 2 if column_spanners else 1
        tbl = doc.add_table(rows=len(df) + num_header_rows, cols=len(df.columns))
        tbl.alignment = WD_TABLE_ALIGNMENT.LEFT
        set_apa_borders(tbl)
        
        current_row = 0
        
        # Column spanners row (if provided)
        if column_spanners:
            spanner_row = tbl.rows[current_row]
            for i in range(len(df.columns)):
                spanner_row.cells[i].text = ""
            
            for spanner_text, start_col, end_col in column_spanners:
                if end_col > start_col:
                    # Merge cells
                    start_cell = spanner_row.cells[start_col]
                    for col in range(start_col + 1, end_col + 1):
                        start_cell.merge(spanner_row.cells[col])
                # Set text and add category-length border
                spanner_row.cells[start_col].text = spanner_text
                spanner_row.cells[start_col].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                spanner_row.cells[start_col].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                for r in spanner_row.cells[start_col].paragraphs[0].runs:
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(11)
                set_cell_border_bottom(spanner_row.cells[start_col])
            current_row += 1
        
        # Header row
        hdr = tbl.rows[current_row]
        for i, col in enumerate(df.columns):
            hdr.cells[i].text = str(col)
            hdr.cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            hdr.cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            for r in hdr.cells[i].paragraphs[0].runs:
                r.font.name = 'Times New Roman'
                r.font.size = Pt(11)
        header_border(hdr)
        current_row += 1
        
        # Data rows
        for ri, (_, row) in enumerate(df.iterrows()):
            for ci, val in enumerate(row):
                cell = tbl.rows[current_row + ri].cells[ci]
                cell.text = str(val) if pd.notna(val) else ''
                cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
                cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT if ci == 0 else WD_ALIGN_PARAGRAPH.RIGHT
                for r in cell.paragraphs[0].runs:
                    r.font.name = 'Times New Roman'
                    r.font.size = Pt(11)
    
    # Note
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    run = p.add_run("Note. ")
    run.italic = True
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    run = p.add_run(note)
    run.font.name = 'Times New Roman'
    run.font.size = Pt(10)
    
    if note_p:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(0)
        run = p.add_run(note_p)
        run.font.name = 'Times New Roman'
        run.font.size = Pt(10)
    
    doc.add_page_break()
    return num + 1


def main():
    parser = argparse.ArgumentParser()
    parser.add_argument('--csv', required=True)
    parser.add_argument('--B', type=int, default=2000)
    parser.add_argument('--ci_type', default='perc')
    parser.add_argument('--out', default=None, help='Output directory (default: same as csv parent)')
    args = parser.parse_args()
    
    csv_path = Path(args.csv)
    out_dir = Path(args.out) if args.out else csv_path.parent
    out_dir.mkdir(parents=True, exist_ok=True)
    if not csv_path.exists():
        print(f"ERROR: {csv_path} not found")
        return
    
    # Handle both CSV and tab-delimited files (.txt from lavaan output)
    df = pd.read_csv(csv_path, sep='\t' if csv_path.suffix == '.txt' else ',')
    
    # Map lavaan column names to expected names
    col_map = {
        'label': 'parameter',
        'ci.lower': 'ci_lower',
        'ci.upper': 'ci_upper'
    }
    df.rename(columns=col_map, inplace=True)
    
    # Handle SE column naming
    if 'se' in df.columns and 'boot_se' not in df.columns:
        df['boot_se'] = df['se']
    
    # Create sig column based on CI excluding zero
    if 'sig' not in df.columns and 'ci_lower' in df.columns and 'ci_upper' in df.columns:
        df['sig'] = (df['ci_lower'] > 0) | (df['ci_upper'] < 0)
    
    B = args.B
    ci = {'bca': 'BCa', 'perc': 'percentile', 'norm': 'normal'}.get(args.ci_type.lower(), args.ci_type)
    
    doc = Document()
    doc.styles['Normal'].font.name = 'Times New Roman'
    doc.styles['Normal'].font.size = Pt(12)
    
    num = 1
    
    # Column spanners for B/SE grouping (columns: Path/Effect=0, B=1, SE=2, 95% CI=3, Sig=4)
    estimate_spanner = [("Estimate", 1, 2)]
    
    # Table 1: Structural paths
    struct = df[df['parameter'].isin(['a1','a1z','a2','a2z','b1','b2','c','cz'])].copy()
    if not struct.empty:
        labels = {'a1':'a₁ (X → M₁)','a1z':'a₁z (X×Z → M₁)','a2':'a₂ (X → M₂)','a2z':'a₂z (X×Z → M₂)',
                  'b1':'b₁ (M₁ → Y)','b2':'b₂ (M₂ → Y)','c':"c′ (X → Y)",'cz':"c′z (X×Z → Y)"}
        struct['Path'] = struct['parameter'].map(labels)
        struct['B'] = struct['est'].apply(lambda x: fmt(x))
        struct['SE'] = struct['boot_se'].apply(lambda x: fmt(x))
        struct['95% CI'] = struct.apply(lambda r: fmt_ci(r['ci_lower'], r['ci_upper']), axis=1)
        struct['Sig'] = struct['sig'].apply(lambda x: '*' if x else '')
        tbl = struct[['Path','B','SE','95% CI','Sig']].reset_index(drop=True)
        num = add_table(doc, num, "Structural Path Coefficients From Bootstrap Analysis", tbl,
                       "X = FASt status; Z = credit dose (centered); M₁ = Emotional Distress; M₂ = Quality of Engagement; Y = Developmental Adjustment.",
                       f"*95% {ci} CI excludes zero. Bootstrap B = {B:,}.",
                       column_spanners=estimate_spanner)
    
    # Table 2: Direct/indirect/total
    rows = []
    for p, lbl in [('c',"Direct effect (c′)"), ('ind_EmoDiss_z_mid','Indirect via EmoDiss'), 
                   ('ind_QualEngag_z_mid','Indirect via QualEngag'), ('total_z_mid','Total effect')]:
        r = df[df['parameter']==p]
        if not r.empty:
            r = r.iloc[0]
            rows.append([lbl, fmt(r['est']), fmt(r['boot_se']), fmt_ci(r['ci_lower'],r['ci_upper']), '*' if r['sig'] else ''])
    if rows:
        tbl = pd.DataFrame(rows, columns=['Effect','B','SE','95% CI','Sig'])
        num = add_table(doc, num, "Direct, Indirect, and Total Effects at Mean Credit Dose", tbl,
                       "Effects at mean credit dose (Z = 0). Indirect = a × b.",
                       f"*95% {ci} CI excludes zero. B = {B:,}.",
                       column_spanners=estimate_spanner)
    
    # Table 3: Conditional indirect + IMM
    rows = []
    for med, med_lbl in [('EmoDiss','EmoDiss'), ('QualEngag','QualEngag')]:
        for lv, lv_lbl in [('low','−1 SD'),('mid','Mean'),('high','+1 SD')]:
            r = df[df['parameter']==f'ind_{med}_z_{lv}']
            if not r.empty:
                r = r.iloc[0]
                rows.append([f'Indirect via {med_lbl} at {lv_lbl}', fmt(r['est']), fmt(r['boot_se']), 
                            fmt_ci(r['ci_lower'],r['ci_upper']), '*' if r['sig'] else ''])
        imm = df[df['parameter']==f'index_MM_{med}']
        if not imm.empty:
            r = imm.iloc[0]
            rows.append([f'IMM ({med_lbl})', fmt(r['est']), fmt(r['boot_se']),
                        fmt_ci(r['ci_lower'],r['ci_upper']), '*' if r['sig'] else ''])
    if rows:
        tbl = pd.DataFrame(rows, columns=['Effect','B','SE','95% CI','Sig'])
        num = add_table(doc, num, "Conditional Indirect Effects and Index of Moderated Mediation", tbl,
                       "Conditional effects at −1 SD, mean, +1 SD of credit dose. IMM = a₁z × b (rate of change in indirect per unit Z).",
                       f"*95% {ci} CI excludes zero. B = {B:,}.",
                       column_spanners=estimate_spanner)
    
    out = out_dir / "Bootstrap_Tables.docx"
    doc.save(str(out))
    print(f"Wrote: {out}")
    print(f"Tables: {num-1} | B = {B:,} | CI = {ci}")


if __name__ == "__main__":
    main()
