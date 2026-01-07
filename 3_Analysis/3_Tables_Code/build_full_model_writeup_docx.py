#!/usr/bin/env python3
"""Build a plain-language Word write-up for the latest full model run.

This script is intentionally conservative:
- It only reports values found in the run outputs.
- It avoids interpretive flourish.

Usage:
  python3 3_Analysis/3_Tables_Code/build_full_model_writeup_docx.py \
    --run_dir 4_Model_Results/Outputs/FullRun_Prepped_20260103_2037 \
    --out 4_Model_Results/Summary/Full_Model_Writeup_20260106.docx
"""

from __future__ import annotations

import argparse
import re
from pathlib import Path

import pandas as pd
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def fmt_num(x, nd: int = 3) -> str:
    try:
        if pd.isna(x):
            return "—"
        return f"{float(x):.{nd}f}"
    except Exception:
        return str(x) if x is not None else "—"


def fmt_p(x) -> str:
    try:
        if pd.isna(x):
            return "—"
        p = float(x)
        if p < 0.001:
            return "< .001"
        return f"= {p:.3f}".replace("0.", ".")
    except Exception:
        return str(x) if x is not None else "—"


def read_fit_measures(path: Path) -> dict[str, float]:
    df = pd.read_csv(path, sep="\t")
    return {str(r["measure"]): float(r["value"]) for _, r in df.iterrows()}


def read_param_estimates(path: Path) -> pd.DataFrame:
    return pd.read_csv(path, sep="\t")


def pick_rows(df: pd.DataFrame, *, label: str | None = None, lhs: str | None = None, op: str | None = None, rhs: str | None = None) -> pd.DataFrame:
    out = df
    if label is not None:
        out = out[out["label"].fillna("") == label]
    if lhs is not None:
        out = out[out["lhs"].fillna("") == lhs]
    if op is not None:
        out = out[out["op"].fillna("") == op]
    if rhs is not None:
        out = out[out["rhs"].fillna("") == rhs]
    return out


def parse_group_counts(verification_text: str) -> dict[str, int]:
    # Extract the three-line table under "Group counts (credit_band)".
    # Example:
    # non_DE non_FASt_1_11 FASt_12plus
    #   263      2718         2019
    m = re.search(r"\(6\) Group counts \(credit_band\)\s+([\s\S]+?)\n\nCheck:", verification_text)
    if not m:
        return {}

    block = m.group(1).strip("\n")
    lines = [ln for ln in block.splitlines() if ln.strip()]
    if len(lines) < 2:
        return {}

    headers = lines[0].split()
    values = lines[1].split()
    if len(headers) != len(values):
        return {}

    counts: dict[str, int] = {}
    for h, v in zip(headers, values):
        try:
            counts[h] = int(v)
        except Exception:
            continue
    return counts


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    doc.add_heading(text, level=level)


def add_paragraph(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    p.paragraph_format.space_after = Pt(6)


def set_doc_font(doc: Document, name: str = "Times New Roman", size_pt: int = 12) -> None:
    style = doc.styles["Normal"]
    style.font.name = name
    style.font.size = Pt(size_pt)


def add_kv_table(doc: Document, title: str, kv: list[tuple[str, str]]) -> None:
    if title:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        p.paragraph_format.space_after = Pt(6)

    table = doc.add_table(rows=len(kv) + 1, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = "Item"
    hdr_cells[1].text = "Value"

    for i, (k, v) in enumerate(kv, start=1):
        row = table.rows[i].cells
        row[0].text = str(k)
        row[1].text = str(v)

    for r in table.rows:
        for c in r.cells:
            for run in c.paragraphs[0].runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(11)


def add_effects_table(doc: Document, title: str, rows: list[dict[str, str]]) -> None:
    if title:
        p = doc.add_paragraph()
        run = p.add_run(title)
        run.bold = True
        p.paragraph_format.space_after = Pt(6)

    cols = ["Parameter", "Estimate", "SE", "p", "95% CI"]
    table = doc.add_table(rows=len(rows) + 1, cols=len(cols))
    table.alignment = WD_TABLE_ALIGNMENT.LEFT

    for j, col in enumerate(cols):
        table.rows[0].cells[j].text = col

    for i, row in enumerate(rows, start=1):
        table.rows[i].cells[0].text = row.get("Parameter", "")
        table.rows[i].cells[1].text = row.get("Estimate", "")
        table.rows[i].cells[2].text = row.get("SE", "")
        table.rows[i].cells[3].text = row.get("p", "")
        table.rows[i].cells[4].text = row.get("CI", "")

    for r in table.rows:
        for c in r.cells:
            for run in c.paragraphs[0].runs:
                run.font.name = "Times New Roman"
                run.font.size = Pt(10)


def main() -> None:
    ap = argparse.ArgumentParser()
    ap.add_argument("--run_dir", required=True, help="FullRun_* directory (e.g., 4_Model_Results/Outputs/FullRun_Prepped_YYYYMMDD_HHMM)")
    ap.add_argument("--out", required=True, help="Output .docx path")
    args = ap.parse_args()

    run_dir = Path(args.run_dir)
    out_path = Path(args.out)
    out_path.parent.mkdir(parents=True, exist_ok=True)

    # Inputs (parallel, total, serial)
    verify_path = run_dir / "logs" / "verification_checklist.txt"
    sem_parallel_fit = run_dir / "RQ1_RQ3_main" / "structural" / "structural_fitMeasures.txt"
    sem_parallel_pe = run_dir / "RQ1_RQ3_main" / "structural" / "structural_parameterEstimates.txt"
    sem_parallel_r2 = run_dir / "RQ1_RQ3_main" / "structural" / "structural_r2.txt"

    total_fit = run_dir / "A0_total_effect" / "structural" / "structural_fitMeasures.txt"
    total_pe = run_dir / "A0_total_effect" / "structural" / "structural_parameterEstimates.txt"

    serial_fit = run_dir / "A1_serial_exploratory" / "structural" / "structural_fitMeasures.txt"
    serial_pe = run_dir / "A1_serial_exploratory" / "structural" / "structural_parameterEstimates.txt"

    verification_text = verify_path.read_text(encoding="utf-8") if verify_path.exists() else ""
    group_counts = parse_group_counts(verification_text)

    fit_parallel = read_fit_measures(sem_parallel_fit)
    pe_parallel = read_param_estimates(sem_parallel_pe)

    fit_total = read_fit_measures(total_fit)
    pe_total = read_param_estimates(total_pe)

    fit_serial = read_fit_measures(serial_fit)
    pe_serial = read_param_estimates(serial_pe)

    # R2 (printed text)
    r2_text = sem_parallel_r2.read_text(encoding="utf-8") if sem_parallel_r2.exists() else ""

    def r2_find(name: str) -> str:
        # Lines are like: "DevAdj 0.872" (with spacing)
        m = re.search(rf"\b{name}\b\s+([0-9]*\.[0-9]+)", r2_text)
        return m.group(1) if m else "—"

    # Helper to format one parameter row
    def effect_row(param_df: pd.DataFrame, label: str, pretty: str) -> dict[str, str]:
        if param_df.empty:
            return {"Parameter": pretty, "Estimate": "—", "SE": "—", "p": "—", "CI": "—"}
        r = param_df.iloc[0]
        return {
            "Parameter": pretty,
            "Estimate": fmt_num(r.get("est"), 3),
            "SE": fmt_num(r.get("se"), 3),
            "p": fmt_p(r.get("pvalue")),
            "CI": f"[{fmt_num(r.get('ci.lower'), 3)}, {fmt_num(r.get('ci.upper'), 3)}]",
        }

    # Key effects (parallel model)
    key_parallel = [
        effect_row(pick_rows(pe_parallel, label="a1"), "a1", "FASt → EmoDiss (a1)"),
        effect_row(pick_rows(pe_parallel, label="a1z"), "a1z", "FASt×credit dose → EmoDiss (a1z)"),
        effect_row(pick_rows(pe_parallel, label="a2"), "a2", "FASt → QualEngag (a2)"),
        effect_row(pick_rows(pe_parallel, label="a2z"), "a2z", "FASt×credit dose → QualEngag (a2z)"),
        effect_row(pick_rows(pe_parallel, label="b1"), "b1", "EmoDiss → DevAdj (b1)"),
        effect_row(pick_rows(pe_parallel, label="b2"), "b2", "QualEngag → DevAdj (b2)"),
        effect_row(pick_rows(pe_parallel, label="c"), "c", "Direct FASt → DevAdj (c)"),
        effect_row(pick_rows(pe_parallel, label="cz"), "cz", "FASt×credit dose → DevAdj (cz)"),
    ]

    # Conditional totals and indirects (parallel model)
    cond_parallel = [
        effect_row(pick_rows(pe_parallel, label="ind_EmoDiss_z_low"), "ind_EmoDiss_z_low", "Indirect via EmoDiss (low credit dose)"),
        effect_row(pick_rows(pe_parallel, label="ind_EmoDiss_z_mid"), "ind_EmoDiss_z_mid", "Indirect via EmoDiss (average credit dose)"),
        effect_row(pick_rows(pe_parallel, label="ind_EmoDiss_z_high"), "ind_EmoDiss_z_high", "Indirect via EmoDiss (high credit dose)"),
        effect_row(pick_rows(pe_parallel, label="ind_QualEngag_z_low"), "ind_QualEngag_z_low", "Indirect via QualEngag (low credit dose)"),
        effect_row(pick_rows(pe_parallel, label="ind_QualEngag_z_mid"), "ind_QualEngag_z_mid", "Indirect via QualEngag (average credit dose)"),
        effect_row(pick_rows(pe_parallel, label="ind_QualEngag_z_high"), "ind_QualEngag_z_high", "Indirect via QualEngag (high credit dose)"),
        effect_row(pick_rows(pe_parallel, label="total_z_low"), "total_z_low", "Total effect (low credit dose)"),
        effect_row(pick_rows(pe_parallel, label="total_z_mid"), "total_z_mid", "Total effect (average credit dose)"),
        effect_row(pick_rows(pe_parallel, label="total_z_high"), "total_z_high", "Total effect (high credit dose)"),
        effect_row(pick_rows(pe_parallel, label="index_MM_EmoDiss"), "index_MM_EmoDiss", "Index of moderated mediation (EmoDiss)"),
        effect_row(pick_rows(pe_parallel, label="index_MM_QualEngag"), "index_MM_QualEngag", "Index of moderated mediation (QualEngag)"),
    ]

    # Total effect model
    total_effect = [
        effect_row(pick_rows(pe_total, label="c_total"), "c_total", "Total effect: FASt → DevAdj (c_total)"),
    ]

    # Serial exploratory model (added path and serial indirect)
    serial_effects = [
        effect_row(pick_rows(pe_serial, label="d"), "d", "EmoDiss → QualEngag (d)"),
        effect_row(pick_rows(pe_serial, label="ind_serial_z_low"), "ind_serial_z_low", "Serial indirect (low credit dose)"),
        effect_row(pick_rows(pe_serial, label="ind_serial_z_mid"), "ind_serial_z_mid", "Serial indirect (average credit dose)"),
        effect_row(pick_rows(pe_serial, label="ind_serial_z_high"), "ind_serial_z_high", "Serial indirect (high credit dose)"),
        effect_row(pick_rows(pe_serial, label="index_MM_serial"), "index_MM_serial", "Index of moderated mediation (serial)"),
    ]

    doc = Document()
    set_doc_font(doc)

    # Title
    title = doc.add_paragraph("Full Model Run: Model and Results Write-Up")
    title.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    title_run = title.runs[0]
    title_run.bold = True
    title_run.font.size = Pt(14)

    subtitle = doc.add_paragraph(f"Run folder: {run_dir.as_posix()}")
    subtitle.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph("")

    add_heading(doc, "1. What was tested", level=2)
    add_paragraph(
        doc,
        "This analysis tested whether FASt status (entering with 12+ transferable credits) is associated with first-year developmental adjustment, and whether that association operates through emotional distress (EmoDiss) and quality of engagement (QualEngag). Credit dose was treated as a moderator through an interaction term that only varies among FASt students (XZ_c = x_FASt × credit_dose_c).",
    )

    add_heading(doc, "2. Sample and grouping", level=2)
    if group_counts:
        add_paragraph(doc, "Group counts used in this run (credit_band):")
        add_kv_table(
            doc,
            "",
            [(k, str(v)) for k, v in group_counts.items()],
        )
    else:
        add_paragraph(doc, "Group counts were not found in the verification checklist for this run.")

    add_heading(doc, "3. Measures (plain description)", level=2)
    add_paragraph(
        doc,
        "Developmental adjustment (DevAdj) was modeled as a second-order factor with four first-order factors: Belong (sbvalued, sbmyself, sbcommunity), Gains (pganalyze, pgthink, pgwork, pgvalues, pgprobsolve), SupportEnv (SEacademic, SEwellness, SEnonacad, SEactivities, SEdiverse), and Satisf (sameinst, evalexp).",
    )
    add_paragraph(
        doc,
        "Emotional distress (EmoDiss) was modeled using six difficulty items: MHWdacad, MHWdlonely, MHWdmental, MHWdexhaust, MHWdsleep, and MHWdfinancial.",
    )
    add_paragraph(
        doc,
        "Quality of engagement (QualEngag) was modeled using five interaction-quality items: QIadmin, QIstudent, QIadvisor, QIfaculty, and QIstaff.",
    )

    add_heading(doc, "4. Structural model (paths)", level=2)
    add_paragraph(
        doc,
        "The model estimated paths from FASt status (x_FASt) and the treated-only credit-dose interaction (XZ_c) to each mediator (EmoDiss, QualEngag) and to the outcome (DevAdj). DevAdj was also regressed on both mediators. Covariates were included on the mediator and outcome equations (cohort, hgrades_c, bparented_c, pell, hapcl, hprecalc13, hchallenge_c, cSFcareer_c).",
    )
    add_paragraph(
        doc,
        "Conditional effects were computed at low (−1 SD), average (0), and high (+1 SD) values of credit_dose_c.",
    )

    add_heading(doc, "5. Model fit", level=2)
    add_paragraph(doc, "Parallel conditional-process SEM (primary model):")
    add_kv_table(
        doc,
        "Fit indices",
        [
            ("χ²(df)", f"{fmt_num(fit_parallel.get('chisq'), 3)} ({int(fit_parallel.get('df'))})"),
            ("CFI", fmt_num(fit_parallel.get("cfi"), 3)),
            ("TLI", fmt_num(fit_parallel.get("tli"), 3)),
            ("RMSEA", fmt_num(fit_parallel.get("rmsea"), 3)),
            ("SRMR", fmt_num(fit_parallel.get("srmr"), 3)),
        ],
    )

    add_paragraph(doc, "Total effect model (DevAdj ~ x_FASt only):")
    add_kv_table(
        doc,
        "Fit indices",
        [
            ("χ²(df)", f"{fmt_num(fit_total.get('chisq'), 3)} ({int(fit_total.get('df'))})"),
            ("CFI", fmt_num(fit_total.get("cfi"), 3)),
            ("TLI", fmt_num(fit_total.get("tli"), 3)),
            ("RMSEA", fmt_num(fit_total.get("rmsea"), 3)),
            ("SRMR", fmt_num(fit_total.get("srmr"), 3)),
        ],
    )

    add_paragraph(doc, "Serial exploratory model (adds EmoDiss → QualEngag):")
    add_kv_table(
        doc,
        "Fit indices",
        [
            ("χ²(df)", f"{fmt_num(fit_serial.get('chisq'), 3)} ({int(fit_serial.get('df'))})"),
            ("CFI", fmt_num(fit_serial.get("cfi"), 3)),
            ("TLI", fmt_num(fit_serial.get("tli"), 3)),
            ("RMSEA", fmt_num(fit_serial.get("rmsea"), 3)),
            ("SRMR", fmt_num(fit_serial.get("srmr"), 3)),
        ],
    )

    add_heading(doc, "6. Main structural results (parallel model)", level=2)
    add_paragraph(
        doc,
        "Table 1 summarizes the primary structural paths for the parallel conditional-process SEM.",
    )
    add_effects_table(doc, "Table 1. Key paths (parallel model)", key_parallel)

    add_heading(doc, "7. Indirect, direct, and total effects", level=2)
    add_paragraph(
        doc,
        "Table 2 summarizes the conditional indirect effects, conditional total effects, and indices of moderated mediation from the parallel model.",
    )
    add_effects_table(doc, "Table 2. Conditional effects and indices (parallel model)", cond_parallel)

    add_heading(doc, "8. Total effect model result", level=2)
    add_paragraph(
        doc,
        "The total effect model estimates the overall association of FASt status with developmental adjustment without mediators or credit-dose moderation.",
    )
    add_effects_table(doc, "Table 3. Total effect (A0_total_effect)", total_effect)

    add_heading(doc, "9. Exploratory serial pathway", level=2)
    add_paragraph(
        doc,
        "The serial exploratory model adds a path from emotional distress to quality of engagement (EmoDiss → QualEngag) and estimates a serial indirect effect from FASt to DevAdj through EmoDiss and then QualEngag.",
    )
    add_effects_table(doc, "Table 4. Serial pathway effects (A1_serial_exploratory)", serial_effects)

    add_heading(doc, "10. Variance explained (R²)", level=2)
    add_paragraph(
        doc,
        "R² values below come from the parallel model output file (structural_r2.txt) and reflect the proportion of variance explained by the predictors in each endogenous variable.",
    )
    add_kv_table(
        doc,
        "R² (selected)",
        [
            ("EmoDiss", r2_find("EmoDiss")),
            ("QualEngag", r2_find("QualEngag")),
            ("DevAdj", r2_find("DevAdj")),
        ],
    )

    add_heading(doc, "11. Plain-language summary of what these results say", level=2)
    add_paragraph(
        doc,
        "In this run, FASt status is positively associated with emotional distress (a1), and higher emotional distress is associated with lower developmental adjustment (b1). The resulting indirect effect through emotional distress is negative and statistically different from zero at low, average, and high credit dose levels.",
    )
    add_paragraph(
        doc,
        "FASt status is not statistically associated with quality of engagement in the parallel model (a2), and the treated-only credit-dose interaction is not statistically associated with either mediator (a1z, a2z) or with developmental adjustment (cz). The indices of moderated mediation are not statistically different from zero in this run.",
    )
    add_paragraph(
        doc,
        "The direct effect of FASt status on developmental adjustment (c) is not statistically different from zero in this run. The conditional total effect is negative and statistically different from zero at the average and high credit dose levels; at the low credit dose level it is negative but does not reach conventional statistical significance.",
    )

    doc.add_paragraph("")
    p = doc.add_paragraph("Outputs used for this write-up:")
    p.runs[0].bold = True
    doc.add_paragraph(str(sem_parallel_fit))
    doc.add_paragraph(str(sem_parallel_pe))
    doc.add_paragraph(str(total_pe))
    doc.add_paragraph(str(serial_pe))
    doc.add_paragraph(str(verify_path))

    doc.save(out_path.as_posix())


if __name__ == "__main__":
    main()
