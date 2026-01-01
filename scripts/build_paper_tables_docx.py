"""Build a single publication-ready Word document of summary tables.

Inputs are read from:
  results/fast_treat_control/official_all_RQs/

Output is written to:
  results/fast_treat_control/official_all_RQs/Paper_Tables_All.docx

Design goals:
- One table per page
- Sequential numbering (Table 1, Table 2, ...)
- Short caption above each table
- APA-ish formatting (readable font size, right-aligned numerics)
- If expected inputs are missing, create a "Missing Output" page listing paths

This script is intentionally robust to small filename/content variations.
"""

from __future__ import annotations

import argparse
import math
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Any, Dict, Iterable, List, Optional, Sequence, Tuple, cast

import pandas as pd
from docx import Document as DocumentFactory
from docx.document import Document as DocxDocument
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Inches, Pt


DEFAULT_BASE_DIR = Path("results/fast_treat_control/official_all_RQs")



# Track which inputs were actually read (for auditability).
USED_INPUT_PATHS: set[Path] = set()


@dataclass
class TableSpec:
    caption: str
    dataframe: Optional[pd.DataFrame]
    missing_paths: List[Path]


def _exists(p: Path) -> bool:
    return p.exists() and p.is_file()


def infer_treatment_var(run_log: Path, rep_data_csv: Path) -> str:
    """Infer the treatment indicator column name.

    Prefers an explicit TREATMENT_VAR line in run_log.txt, otherwise inspects
    the rep_data CSV for known column names.
    """
    if _exists(run_log):
        txt = read_text_any(run_log)
        m = re.search(r"^TREATMENT_VAR:\s*(\S+)", txt, flags=re.MULTILINE)
        if m:
            return m.group(1).strip()

    if _exists(rep_data_csv):
        try:
            d = pd.read_csv(rep_data_csv, nrows=5)
            cols = set(d.columns)
            for cand in ["x_FASt", "x_DE"]:
                if cand in cols:
                    return cand
        except Exception:
            pass

    # Conservative fallback
    return "x_FASt"


def infer_qualengage_col(df_columns: Iterable[str]) -> Optional[str]:
    cols = set(df_columns)
    for cand in ["QualEngage", "QualEngag"]:
        if cand in cols:
            return cand
    return None


def read_text_any(path: Path) -> str:
    USED_INPUT_PATHS.add(path)
    return path.read_text(encoding="utf-8", errors="ignore")


def read_table_any(path: Path) -> pd.DataFrame:
    """Read a table from a .txt/.tsv-like file.

    Prefers tab-delimited, falls back to whitespace.
    """
    USED_INPUT_PATHS.add(path)
    try:
        return cast(pd.DataFrame, pd.read_csv(str(path), sep="\t", dtype=str, keep_default_na=False))
    except Exception:
        # sep regex avoids the deprecated/poorly-typed delim_whitespace signature
        return cast(
            pd.DataFrame,
            pd.read_csv(str(path), sep=r"\s+", engine="python", dtype=str, keep_default_na=False),
        )


def as_numeric(series: pd.Series) -> pd.Series:
    # Pandas' typing allows scalar returns; we only call this with Series.
    return cast(pd.Series, pd.to_numeric(series, errors="coerce"))


def _parse_float(x: object) -> Optional[float]:
    if x is None:
        return None
    if isinstance(x, str):
        s = x.strip()
        if s == "":
            return None
        try:
            return float(s)
        except Exception:
            return None
    try:
        return float(str(x).strip())
    except Exception:
        return None


def fmt_num(x: object, nd: int = 3) -> str:
    if x is None:
        return ""
    if isinstance(x, str) and x.strip() == "":
        return ""
    v = _parse_float(x)
    if v is None:
        return str(x)

    if math.isnan(v):
        return ""
    if abs(v) >= 1000:
        return f"{v:,.{nd}f}"
    return f"{v:.{nd}f}"


def fmt_p(x: object) -> str:
    v = _parse_float(x)
    if v is None:
        return ""
    if math.isnan(v):
        return ""
    if v < 0.001:
        return "< .001"
    s = f"{v:.3f}"
    if s.startswith("0"):
        s = s[1:]
    return s


def fmt_ci(lo: object, hi: object, nd: int = 3) -> str:
    lo_s = fmt_num(lo, nd=nd)
    hi_s = fmt_num(hi, nd=nd)
    if lo_s == "" and hi_s == "":
        return ""
    return f"[{lo_s}, {hi_s}]"


def normalize_columns(df: pd.DataFrame) -> pd.DataFrame:
    df = df.copy()
    df.columns = [str(c).strip() for c in df.columns]
    return df


def _parse_ps_model_line(psw_stage_txt: str) -> Optional[str]:
    m = re.search(r"PS model:\s*(.+)", psw_stage_txt)
    if m:
        return m.group(1).strip()
    return None


def _parse_ps_covariates(ps_model_line: str) -> List[str]:
    # Example: x_DE ~ hgrades_c + bparented_c + ...
    if "~" not in ps_model_line:
        return []
    rhs = ps_model_line.split("~", 1)[1]
    covs = [c.strip() for c in rhs.split("+")]
    return [c for c in covs if c]


def table_sample_split(rep_data_csv: Path, treatment_var: str) -> TableSpec:
    """Analytic sample size and treatment split."""
    missing: List[Path] = []
    if not _exists(rep_data_csv):
        missing.append(rep_data_csv)
        return TableSpec(
            caption=(
                "Analytic sample and treatment split. Note. Estimand targets the overlap population (ATO) via PSW overlap weights."
            ),
            dataframe=None,
            missing_paths=missing,
        )

    d = pd.read_csv(rep_data_csv)
    if treatment_var not in d.columns:
        missing.append(rep_data_csv)
        return TableSpec(
            caption=(
                "Analytic sample and treatment split. Note. Estimand targets the overlap population (ATO) via PSW overlap weights."
            ),
            dataframe=None,
            missing_paths=missing,
        )

    x = pd.to_numeric(d[treatment_var], errors="coerce")
    n_total = int(x.notna().sum())
    n_treat = int((x == 1).sum())
    n_control = int((x == 0).sum())
    pct_treat = (n_treat / n_total * 100.0) if n_total > 0 else None

    df = pd.DataFrame(
        [
            ("N (analytic)", f"{n_total:,d}"),
            (f"Treatment: {treatment_var} = 1", f"{n_treat:,d} ({pct_treat:.1f}%)" if pct_treat is not None else f"{n_treat:,d}"),
            (f"Control: {treatment_var} = 0", f"{n_control:,d}"),
        ],
        columns=["Field", "Value"],
    )

    return TableSpec(
        caption="Analytic sample and treatment split. Note. Estimand targets the overlap population (ATO) via PSW overlap weights.",
        dataframe=df,
        missing_paths=[],
    )


def table_run_log(run_log: Path, psw_stage: Optional[Path]) -> TableSpec:
    missing: List[Path] = []
    rows: List[Tuple[str, str]] = []

    if not _exists(run_log):
        missing.append(run_log)
    else:
        text = read_text_any(run_log).splitlines()
        for line in text:
            if ":" in line:
                k, v = line.split(":", 1)
                k = k.strip()
                v = v.strip()
                if k and v:
                    rows.append((k, v))

    # If the file exists but contains no usable rows, treat it as a failure.
    if _exists(run_log) and not rows:
        missing.append(run_log)

    # Add a short PSW model line if present
    if psw_stage is not None:
        if not _exists(psw_stage):
            missing.append(psw_stage)
        else:
            txt = read_text_any(psw_stage)
            m = re.search(r"PS model:\s*(.+)", txt)
            if m:
                rows.append(("PS model", m.group(1).strip()))

    df = pd.DataFrame(rows, columns=["Field", "Value"]) if rows else None
    return TableSpec(
        caption="Sample/design summary (run metadata and PSW model).",
        dataframe=df,
        missing_paths=missing,
    )


def compute_weight_diagnostics(psw_stage: Path, rep_data_csv: Path) -> TableSpec:
    missing: List[Path] = []

    if not _exists(psw_stage):
        missing.append(psw_stage)
        return TableSpec(
            caption="Weight diagnostics (distribution + effective sample size).",
            dataframe=None,
            missing_paths=missing,
        )

    txt = read_text_any(psw_stage)

    # Parse the 6-number summary block if present (Min, 1st Qu., Median, Mean, 3rd Qu., Max.)
    summary_vals = {}
    summary_match = re.search(
        r"Weights summary \(non-missing\):\s*\n\s*Min\.[^\n]*\n\s*([0-9eE+\-.]+)\s+([0-9eE+\-.]+)\s+([0-9eE+\-.]+)\s+([0-9eE+\-.]+)\s+([0-9eE+\-.]+)\s+([0-9eE+\-.]+)",
        txt,
        flags=re.MULTILINE,
    )
    if summary_match:
        summary_vals = {
            "Min": summary_match.group(1),
            "Q1": summary_match.group(2),
            "Median": summary_match.group(3),
            "Mean": summary_match.group(4),
            "Q3": summary_match.group(5),
            "Max": summary_match.group(6),
        }

    ess = None
    if not _exists(rep_data_csv):
        missing.append(rep_data_csv)
    else:
        d = pd.read_csv(rep_data_csv)
        # weight column could be psw
        wcol = "psw" if "psw" in d.columns else None
        if wcol is not None:
            w_series = cast(pd.Series, pd.to_numeric(d[wcol], errors="coerce"))
            w = w_series.dropna()
            if len(w) > 0:
                ess = (w.sum() ** 2) / (w.pow(2).sum())

    rows = []
    for k in ["Min", "Q1", "Median", "Mean", "Q3", "Max"]:
        if k in summary_vals:
            rows.append((k, fmt_num(summary_vals[k], nd=3)))
    if ess is not None:
        rows.append(("ESS", f"{ess:,.1f}"))

    df = pd.DataFrame(rows, columns=["Diagnostic", "Value"]) if rows else None

    # If expected source exists but produces no usable output, treat as failure.
    if _exists(psw_stage) and (df is None or df.empty):
        missing.append(psw_stage)

    return TableSpec(
        caption="PSW overlap weight diagnostics (distribution and effective sample size).",
        dataframe=df,
        missing_paths=missing,
    )


def table_balance(balance_path: Path) -> TableSpec:
    missing: List[Path] = []
    if not _exists(balance_path):
        missing.append(balance_path)
        return TableSpec(
            caption="Covariate balance (standardized mean differences; unweighted vs PSW-weighted).",
            dataframe=None,
            missing_paths=missing,
        )

    df = read_table_any(balance_path)
    df = normalize_columns(df)
    # IMPORTANT: preserve exact SMD strings from the file (committee request).

    if df.empty:
        missing.append(balance_path)
        return TableSpec(
            caption="Covariate balance after PSW overlap weights (SMDs).",
            dataframe=None,
            missing_paths=missing,
        )
    return TableSpec(
        caption="Covariate balance after PSW overlap weights (SMDs).",
        dataframe=df,
        missing_paths=missing,
    )


def table_fit_indices(fit_path: Path, caption: str) -> TableSpec:
    missing: List[Path] = []
    if not _exists(fit_path):
        missing.append(fit_path)
        return TableSpec(caption=caption, dataframe=None, missing_paths=missing)

    df = read_table_any(fit_path)
    df = normalize_columns(df)
    if set(df.columns) >= {"measure", "value"}:
        df = df[["measure", "value"]]
        # keep common indices only, in a nicer order when possible
        keep = [
            "df",
            "chisq",
            "pvalue",
            "cfi",
            "tli",
            "rmsea",
            "srmr",
            "cfi.scaled",
            "tli.scaled",
            "rmsea.scaled",
            "cfi.robust",
            "tli.robust",
            "rmsea.robust",
        ]
        df["measure"] = df["measure"].astype(str)
        df["value"] = df["value"].astype(str)
        df = df[df["measure"].isin(keep)].copy()
        df["_ord"] = df["measure"].apply(lambda m: keep.index(m) if m in keep else 999)
        df = df.sort_values(["_ord"]).drop(columns=["_ord"])
        df.rename(columns={"measure": "Index", "value": "Value"}, inplace=True)
        df["Value"] = as_numeric(df["Value"]).map(lambda v: fmt_num(v, nd=3) if pd.notna(v) else "")

    if df is None or df.empty:
        missing.append(fit_path)
        return TableSpec(caption=caption, dataframe=None, missing_paths=missing)

    return TableSpec(caption=caption, dataframe=df, missing_paths=missing)


def _extract_paths(df: pd.DataFrame, include_labels: Optional[Sequence[str]] = None) -> pd.DataFrame:
    df = normalize_columns(df)
    needed_cols = {"lhs", "op", "rhs"}
    if not needed_cols.issubset(set(df.columns)):
        return pd.DataFrame()

    d = df.copy()
    d = d[d["op"].astype(str) == "~"].copy()

    if "label" in d.columns:
        d["label"] = d["label"].astype(str)
        d = d[d["label"].str.strip() != ""].copy()
        if include_labels is not None:
            d = d[d["label"].isin(set(include_labels))].copy()

    d["Path"] = d["lhs"].astype(str) + " ~ " + d["rhs"].astype(str)

    out_cols = ["Path"]
    if "label" in d.columns:
        out_cols.append("label")

    for col in ["est", "se", "z", "pvalue", "ci.lower", "ci.upper", "std.all"]:
        if col in d.columns:
            out_cols.append(col)

    # Use .loc to ensure a DataFrame (even if out_cols changes).
    d = d.loc[:, cast(List[str], out_cols)].copy()

    # Friendly renames
    rename = {
        "label": "Label",
        "est": "Estimate",
        "se": "SE",
        "z": "z",
        "pvalue": "p",
        "std.all": "Std (all)",
    }
    d = d.rename(mapper=rename, axis=1)

    if "p" in d.columns:
        d["p"] = as_numeric(d["p"]).map(fmt_p)
    if "Estimate" in d.columns:
        d["Estimate"] = as_numeric(d["Estimate"]).map(lambda v: fmt_num(v, nd=3) if pd.notna(v) else "")
    if "SE" in d.columns:
        d["SE"] = as_numeric(d["SE"]).map(lambda v: fmt_num(v, nd=3) if pd.notna(v) else "")
    if "z" in d.columns:
        d["z"] = as_numeric(d["z"]).map(lambda v: fmt_num(v, nd=3) if pd.notna(v) else "")
    if "Std (all)" in d.columns:
        d["Std (all)"] = as_numeric(d["Std (all)"].map(lambda v: v)).map(
            lambda v: fmt_num(v, nd=3) if pd.notna(v) else ""
        )

    if "ci.lower" in d.columns and "ci.upper" in d.columns:
        lo = d["ci.lower"].copy()
        hi = d["ci.upper"].copy()
        d["95% CI"] = [fmt_ci(a, b, nd=3) for a, b in zip(lo, hi)]
        d.drop(columns=["ci.lower", "ci.upper"], inplace=True)

    return d


def _extract_defined(df: pd.DataFrame) -> pd.DataFrame:
    df = normalize_columns(df)
    if not {"lhs", "op"}.issubset(df.columns):
        return pd.DataFrame()
    d: pd.DataFrame = df.copy()
    d = cast(pd.DataFrame, d[d["op"].astype(str) == ":="].copy())
    if d.empty:
        return cast(pd.DataFrame, d)

    # Avoid pandas.rename typing-stub issues; do a deterministic column remap.
    rename_map = {
        "lhs": "Defined parameter",
        "est": "Estimate",
        "se": "SE",
        "z": "z",
        "pvalue": "p",
        "std.all": "Std (all)",
    }
    d.columns = [rename_map.get(str(c), str(c)) for c in d.columns]
    keep = [c for c in ["Defined parameter", "Estimate", "SE", "z", "p", "ci.lower", "ci.upper"] if c in d.columns]
    d = d.loc[:, keep].copy()

    if "p" in d.columns:
        d["p"] = as_numeric(d["p"]).map(fmt_p)
    if "Estimate" in d.columns:
        d["Estimate"] = as_numeric(d["Estimate"]).map(lambda v: fmt_num(v, nd=3) if pd.notna(v) else "")
    if "SE" in d.columns:
        d["SE"] = as_numeric(d["SE"]).map(lambda v: fmt_num(v, nd=3) if pd.notna(v) else "")
    if "z" in d.columns:
        d["z"] = as_numeric(d["z"]).map(lambda v: fmt_num(v, nd=3) if pd.notna(v) else "")

    if "ci.lower" in d.columns and "ci.upper" in d.columns:
        d["95% CI"] = [fmt_ci(a, b, nd=3) for a, b in zip(d["ci.lower"], d["ci.upper"])]
        d.drop(columns=["ci.lower", "ci.upper"], inplace=True)

    return cast(pd.DataFrame, d)


def table_structural_paths(pe_path: Path, caption: str, include_labels: Optional[Sequence[str]] = None) -> TableSpec:
    missing: List[Path] = []
    if not _exists(pe_path):
        missing.append(pe_path)
        return TableSpec(caption=caption, dataframe=None, missing_paths=missing)

    df = cast(pd.DataFrame, read_table_any(pe_path))
    d = _extract_paths(df, include_labels=include_labels)
    if d is None or d.empty:
        missing.append(pe_path)
        return TableSpec(caption=caption, dataframe=None, missing_paths=missing)
    return TableSpec(caption=caption, dataframe=d, missing_paths=missing)


def table_defined_params(pe_path: Path, caption: str) -> TableSpec:
    missing: List[Path] = []
    if not _exists(pe_path):
        missing.append(pe_path)
        return TableSpec(caption=caption, dataframe=None, missing_paths=missing)

    df = cast(pd.DataFrame, read_table_any(pe_path))
    d = _extract_defined(df)
    if d is None or d.empty:
        missing.append(pe_path)
        return TableSpec(caption=caption, dataframe=None, missing_paths=missing)
    return TableSpec(caption=caption, dataframe=d, missing_paths=missing)


def table_standardized_loadings(stdsol_path: Path) -> TableSpec:
    missing: List[Path] = []
    if not _exists(stdsol_path):
        missing.append(stdsol_path)
        return TableSpec(
            caption="Standardized factor loadings (primary pooled model).",
            dataframe=None,
            missing_paths=missing,
        )

    df = cast(pd.DataFrame, read_table_any(stdsol_path))
    df = normalize_columns(df)
    if not {"lhs", "op", "rhs"}.issubset(df.columns):
        return TableSpec(
            caption="Standardized factor loadings (primary pooled model).",
            dataframe=None,
            missing_paths=[stdsol_path],
        )

    d = df[df["op"].astype(str) == "=~"].copy()
    if d.empty:
        return TableSpec(
            caption="Standardized factor loadings (primary pooled model).",
            dataframe=None,
            missing_paths=[stdsol_path],
        )

    out = pd.DataFrame(
        {
            "Factor": d["lhs"].astype(str),
            "Indicator": d["rhs"].astype(str),
        }
    )
    if "est.std" in d.columns:
        out["Std. loading"] = as_numeric(d["est.std"]).map(lambda v: fmt_num(v, nd=3) if pd.notna(v) else "")
    elif "std.all" in d.columns:
        out["Std. loading"] = as_numeric(d["std.all"]).map(lambda v: fmt_num(v, nd=3) if pd.notna(v) else "")

    return TableSpec(
        caption="Standardized factor loadings (primary pooled model).",
        dataframe=out,
        missing_paths=[],
    )


def table_r2_summary(r2_path: Path) -> TableSpec:
    missing: List[Path] = []
    if not _exists(r2_path):
        missing.append(r2_path)
        return TableSpec(
            caption="Structural model R² for key endogenous variables (weighted pooled model).",
            dataframe=None,
            missing_paths=missing,
        )

    txt = read_text_any(r2_path)
    # Parse a simple name/value space-separated format
    pairs = re.findall(r"\b([A-Za-z0-9_]+)\s+([0-9.]+)", txt)
    r2_map: Dict[str, str] = {k: v for k, v in pairs}
    keep = ["DevAdj", "EmoDiss", "QualEngag"]
    rows: List[Tuple[str, str]] = []
    for k in keep:
        if k in r2_map:
            rows.append((k, fmt_num(r2_map[k], nd=3)))

    df = pd.DataFrame(rows, columns=["Outcome", "R²"]) if rows else None
    return TableSpec(
        caption="Structural model R² for key endogenous variables (weighted pooled model).",
        dataframe=df,
        missing_paths=[r2_path] if (df is None or df.empty) else [],
    )


def table_invariance_summary(meas_root: Path, w_vars: Sequence[str]) -> TableSpec:
    """Single stacked table: one row per W x model step, with Δfit where applicable."""
    missing: List[Path] = []
    rows: List[Dict[str, str]] = []

    for w in w_vars:
        by_dir = meas_root / f"by_{w}"
        deltas_path = by_dir / "fit_change_deltas.txt"
        stack_path = by_dir / "fit_index_stack.txt"

        if not by_dir.exists():
            missing.append(by_dir)
            continue

        df_deltas = read_table_any(deltas_path) if _exists(deltas_path) else None
        if df_deltas is None:
            missing.append(deltas_path)
        else:
            df_deltas = normalize_columns(df_deltas)

        df_stack = read_table_any(stack_path) if _exists(stack_path) else None
        if df_stack is None:
            missing.append(stack_path)
            continue

        df_stack = normalize_columns(df_stack)
        if not {"model", "measure", "value"}.issubset(df_stack.columns):
            missing.append(stack_path)
            continue

        wide = cast(
            pd.DataFrame,
            df_stack.pivot_table(index="model", columns="measure", values="value", aggfunc="first").reset_index(),
        )

        def _get_delta(step: str, col: str) -> str:
            if df_deltas is None or df_deltas.empty:
                return ""
            if "step" not in df_deltas.columns or col not in df_deltas.columns:
                return ""
            hit = df_deltas[df_deltas["step"].astype(str) == step]
            if hit.empty:
                return ""
            return fmt_num(hit.iloc[0][col], nd=3)

        # order models as produced by the pipeline
        model_order = ["configural", "metric_firstorder", "metric_secondorder", "scalar"]
        for m in model_order:
            hit = wide[wide["model"].astype(str) == m]
            if hit.empty:
                continue
            row0 = hit.iloc[0].to_dict()
            cfi = fmt_num(row0.get("cfi.robust", row0.get("cfi", "")), nd=3)
            rmsea = fmt_num(row0.get("rmsea.robust", row0.get("rmsea", "")), nd=3)
            srmr = fmt_num(row0.get("srmr", ""), nd=3)

            delta_cfi = ""
            delta_rmsea = ""
            delta_srmr = ""
            if m == "metric_firstorder":
                delta_cfi = _get_delta("configural_to_metric_firstorder", "delta_cfi")
                delta_rmsea = _get_delta("configural_to_metric_firstorder", "delta_rmsea")
                delta_srmr = _get_delta("configural_to_metric_firstorder", "delta_srmr")
            elif m == "metric_secondorder":
                delta_cfi = _get_delta("metric_firstorder_to_metric_secondorder", "delta_cfi")
                delta_rmsea = _get_delta("metric_firstorder_to_metric_secondorder", "delta_rmsea")
                delta_srmr = _get_delta("metric_firstorder_to_metric_secondorder", "delta_srmr")
            elif m == "scalar":
                delta_cfi = _get_delta("metric_secondorder_to_scalar", "delta_cfi")
                delta_rmsea = _get_delta("metric_secondorder_to_scalar", "delta_rmsea")
                delta_srmr = _get_delta("metric_secondorder_to_scalar", "delta_srmr")

            rows.append(
                {
                    "W": w,
                    "Model": m,
                    "CFI (robust)": cfi,
                    "RMSEA (robust)": rmsea,
                    "SRMR": srmr,
                    "ΔCFI": delta_cfi,
                    "ΔRMSEA": delta_rmsea,
                    "ΔSRMR": delta_srmr,
                }
            )

    df_out = pd.DataFrame(rows) if rows else None
    cap = (
        "Measurement invariance by W (configural → metric → scalar). "
        "Note. Common guidelines: ΔCFI ≤ .010 and ΔRMSEA ≤ .015 (Cheung & Rensvold, 2002; Chen, 2007)."
    )
    return TableSpec(caption=cap, dataframe=df_out, missing_paths=missing if (df_out is None or df_out.empty) else [])


def table_race_fit_indices(race_root: Path) -> TableSpec:
    missing: List[Path] = []
    rows: List[Dict[str, str]] = []
    if not race_root.exists():
        return TableSpec(
            caption="Race-disaggregated structural model fit indices.",
            dataframe=None,
            missing_paths=[race_root],
        )

    for race_dir in sorted([p for p in race_root.iterdir() if p.is_dir()]):
        fit_path = race_dir / "structural" / "structural_fitMeasures.txt"
        if not _exists(fit_path):
            missing.append(fit_path)
            continue
        df = normalize_columns(read_table_any(fit_path))
        if not {"measure", "value"}.issubset(df.columns):
            missing.append(fit_path)
            continue
        d = dict(zip(df["measure"].astype(str), df["value"].astype(str)))
        rows.append(
            {
                "Race": race_dir.name,
                "CFI (robust)": fmt_num(d.get("cfi.robust", d.get("cfi", "")), nd=3),
                "RMSEA (robust)": fmt_num(d.get("rmsea.robust", d.get("rmsea", "")), nd=3),
                "SRMR": fmt_num(d.get("srmr", ""), nd=3),
            }
        )

    df_out = pd.DataFrame(rows) if rows else None
    return TableSpec(
        caption="Race-disaggregated structural model fit indices. Note. Groups with n below MIN_RACE_N are omitted.",
        dataframe=df_out,
        missing_paths=missing if (df_out is None or df_out.empty) else [],
    )


def table_race_key_paths(race_root: Path, key_labels: Sequence[str]) -> TableSpec:
    missing: List[Path] = []
    rows: List[pd.DataFrame] = []

    if not race_root.exists():
        return TableSpec(
            caption="Race-disaggregated key coefficients (by race).",
            dataframe=None,
            missing_paths=[race_root],
        )

    for race_dir in sorted([p for p in race_root.iterdir() if p.is_dir()]):
        pe_path = race_dir / "structural" / "structural_parameterEstimates.txt"
        if not _exists(pe_path):
            missing.append(pe_path)
            continue
        df = normalize_columns(read_table_any(pe_path))
        if "label" not in df.columns:
            missing.append(pe_path)
            continue
        df = df[df["label"].astype(str).isin(set(key_labels))].copy()
        if df.empty:
            continue
        extracted = _extract_paths(df)
        if extracted.empty:
            continue
        extracted.insert(0, "Race", race_dir.name)
        rows.append(extracted)

    df_out = pd.concat(rows, ignore_index=True) if rows else None
    return TableSpec(
        caption="Race-disaggregated key coefficients (by race). Note. Groups with n below MIN_RACE_N are omitted.",
        dataframe=df_out,
        missing_paths=missing if (df_out is None or df_out.empty) else [],
    )


def table_descriptives_by_treatment(rep_data_csv: Path, treatment_var: str) -> TableSpec:
    missing: List[Path] = []
    if not _exists(rep_data_csv):
        missing.append(rep_data_csv)
        return TableSpec(
            caption="Descriptives by treatment (unweighted and PSW-weighted).",
            dataframe=None,
            missing_paths=missing,
        )

    d = pd.read_csv(rep_data_csv)
    qual_col = infer_qualengage_col(d.columns)
    needed = {treatment_var, "DevAdj", "EmoDiss", "trnsfr_cr"}
    if qual_col is None:
        missing.append(rep_data_csv)
        return TableSpec(
            caption="Descriptives by treatment (unweighted and PSW-weighted).",
            dataframe=None,
            missing_paths=missing,
        )
    needed.add(qual_col)

    has_psw = "psw" in d.columns
    if has_psw:
        needed.add("psw")

    if not needed.issubset(set(d.columns)):
        missing.append(rep_data_csv)
        return TableSpec(
            caption="Descriptives by treatment (unweighted and PSW-weighted).",
            dataframe=None,
            missing_paths=missing,
        )

    def w_mean(x: pd.Series, w: pd.Series) -> float:
        ww = w.copy()
        xx = x.copy()
        m = ww.notna() & xx.notna()
        ww = ww[m]
        xx = xx[m]
        return float((ww * xx).sum() / ww.sum()) if ww.sum() != 0 else float("nan")

    def w_sd(x: pd.Series, w: pd.Series) -> float:
        mu = w_mean(x, w)
        ww = w.copy()
        xx = x.copy()
        m = ww.notna() & xx.notna()
        ww = ww[m]
        xx = xx[m]
        return float(math.sqrt(((ww * (xx - mu) ** 2).sum() / ww.sum()))) if ww.sum() != 0 else float("nan")

    x = pd.to_numeric(d[treatment_var], errors="coerce")
    w = pd.to_numeric(d["psw"], errors="coerce") if has_psw else pd.Series([float("nan")] * len(d))

    variables = [
        ("DevAdj", "DevAdj"),
        ("EmoDiss", "EmoDiss"),
        ("QualEngag", cast(str, qual_col)),
        ("Transfer credits at entry", "trnsfr_cr"),
    ]
    rows: List[Dict[str, str]] = []
    for label, col in variables:
        v = pd.to_numeric(d[col], errors="coerce")
        for g, gname in [(0, "Control"), (1, "Treatment")]:
            m = x == g
            un_m = float(v[m].mean())
            un_sd = float(v[m].std(ddof=1))
            wt_m = w_mean(v[m], w[m]) if has_psw else float("nan")
            wt_sd = w_sd(v[m], w[m]) if has_psw else float("nan")
            rows.append(
                {
                    "Variable": label,
                    "Group": gname,
                    "Unweighted M (SD)": f"{un_m:.3f} ({un_sd:.3f})" if not math.isnan(un_m) else "",
                    "PSW-weighted M (SD)": f"{wt_m:.3f} ({wt_sd:.3f})" if not math.isnan(wt_m) else ("(no psw column)" if not has_psw else ""),
                }
            )

    df_out = pd.DataFrame(rows) if rows else None
    return TableSpec(
        caption="Descriptives by treatment (unweighted and PSW-weighted).",
        dataframe=df_out,
        missing_paths=missing if (df_out is None or df_out.empty) else [],
    )


def table_sensitivity_unweighted_vs_weighted(weighted_pe: Path, unweighted_pe: Path, key_labels: Sequence[str]) -> TableSpec:
    missing: List[Path] = []
    if not _exists(weighted_pe):
        missing.append(weighted_pe)
    if not _exists(unweighted_pe):
        missing.append(unweighted_pe)
    if missing:
        return TableSpec(
            caption="Sensitivity/spec checks: unweighted vs PSW-weighted headline paths.",
            dataframe=None,
            missing_paths=missing,
        )

    wdf = normalize_columns(read_table_any(weighted_pe))
    udf = normalize_columns(read_table_any(unweighted_pe))

    if "label" not in wdf.columns or "label" not in udf.columns:
        return TableSpec(
            caption="Sensitivity/spec checks: unweighted vs PSW-weighted headline paths.",
            dataframe=None,
            missing_paths=[weighted_pe, unweighted_pe],
        )

    wsub = _extract_paths(wdf[wdf["label"].astype(str).isin(set(key_labels))].copy())
    usub = _extract_paths(udf[udf["label"].astype(str).isin(set(key_labels))].copy())

    # Rename CI label for the weighted model (bca.simple)
    if "95% CI" in wsub.columns:
        wsub = wsub.rename(columns={"95% CI": "Bias-corrected bootstrap CI (bca.simple)"})

    # Merge by Label (stable identifier)
    if "Label" not in wsub.columns or "Label" not in usub.columns:
        return TableSpec(
            caption="Sensitivity/spec checks: unweighted vs PSW-weighted headline paths.",
            dataframe=None,
            missing_paths=[weighted_pe, unweighted_pe],
        )

    merged = pd.merge(
        wsub,
        usub,
        on=["Label"],
        how="outer",
        suffixes=(" (PSW)", " (Unweighted)"),
    )
    # Keep a compact column set when available
    keep_cols = [
        c
        for c in [
            "Path (PSW)",
            "Label",
            "Estimate (PSW)",
            "Bias-corrected bootstrap CI (bca.simple)",
            "Estimate (Unweighted)",
            "SE (Unweighted)",
            "p (Unweighted)",
        ]
        if c in merged.columns
    ]
    out = merged[keep_cols].copy() if keep_cols else merged

    return TableSpec(
        caption="Sensitivity/spec checks: unweighted vs PSW-weighted headline paths. Note. Unweighted model is fit on the same analytic sample with no weight trimming.",
        dataframe=out,
        missing_paths=[],
    )


def table_total_effect(pe_path: Path) -> TableSpec:
    # try to extract c_total row(s)
    missing: List[Path] = []
    if not _exists(pe_path):
        missing.append(pe_path)
        return TableSpec(
            caption="Total effect model (Eq. 1): DevAdj ~ treatment (total effect).",
            dataframe=None,
            missing_paths=missing,
        )

    df = cast(pd.DataFrame, read_table_any(pe_path))
    df = normalize_columns(df)
    if "label" in df.columns:
        df = cast(pd.DataFrame, df[df["label"].astype(str).isin({"c_total"})].copy())
    d = _extract_paths(cast(pd.DataFrame, df))
    if d is None or d.empty:
        missing.append(pe_path)
        return TableSpec(
            caption="Total effect (Eq. 1): estimated total effect of treatment on DevAdj.",
            dataframe=None,
            missing_paths=missing,
        )
    return TableSpec(
        caption="Total effect (Eq. 1): estimated total effect of treatment on DevAdj.",
        dataframe=d,
        missing_paths=missing,
    )


def table_invariance_for_W(by_w_dir: Path, w_name: str) -> List[TableSpec]:
    specs: List[TableSpec] = []

    deltas = by_w_dir / "fit_change_deltas.txt"
    stack = by_w_dir / "fit_index_stack.txt"

    missing: List[Path] = []
    df_deltas: Optional[pd.DataFrame] = None
    df_stack: Optional[pd.DataFrame] = None

    if _exists(deltas):
        df_deltas = cast(pd.DataFrame, read_table_any(deltas))
        df_deltas = normalize_columns(df_deltas)
        for col in ["delta_cfi", "delta_rmsea", "delta_srmr"]:
            if col in df_deltas.columns:
                df_deltas[col] = as_numeric(df_deltas[col]).map(lambda v: fmt_num(v, nd=3) if pd.notna(v) else "")
    else:
        missing.append(deltas)

    if _exists(stack):
        df_stack = cast(pd.DataFrame, read_table_any(stack))
        df_stack = normalize_columns(df_stack)
        # reshape to wide: one row per model
        if set(df_stack.columns) >= {"model", "measure", "value"}:
            wide = cast(
                pd.DataFrame,
                df_stack.pivot_table(index="model", columns="measure", values="value", aggfunc="first").reset_index(),
            )
            # keep common indices
            keep = ["model", "df", "chisq", "cfi", "tli", "rmsea", "srmr"]
            cols = [c for c in keep if c in wide.columns]
            wide = wide[cols].copy()
            for col in ["df", "chisq", "cfi", "tli", "rmsea", "srmr"]:
                if col in wide.columns:
                    wide[col] = as_numeric(wide[col]).map(lambda v: fmt_num(v, nd=3) if pd.notna(v) else "")
            # Avoid pandas.rename typing-stub issues
            wide.columns = ["Model" if str(c) == "model" else str(c) for c in wide.columns]
            df_stack = cast(pd.DataFrame, wide)
    else:
        missing.append(stack)

    specs.append(
        TableSpec(
            caption=f"Measurement invariance by {w_name}: fit indices (configural/metric/scalar).",
            dataframe=df_stack if df_stack is not None and not df_stack.empty else None,
            missing_paths=[stack] if (stack in missing or df_stack is None or (hasattr(df_stack, "empty") and df_stack.empty)) else [],
        )
    )
    specs.append(
        TableSpec(
            caption=f"Measurement invariance by {w_name}: Δfit between nested models.",
            dataframe=df_deltas if df_deltas is not None and not df_deltas.empty else None,
            missing_paths=[deltas] if (deltas in missing or df_deltas is None or (hasattr(df_deltas, "empty") and df_deltas.empty)) else [],
        )
    )

    return specs


def table_mg_structural_for_W(w_dir: Path, w_label: str) -> List[TableSpec]:
    """Group-specific key paths and defined-parameter contrasts for one W."""
    specs: List[TableSpec] = []

    pe_path = w_dir / "structural" / "structural_parameterEstimates.txt"
    ref_path = w_dir / "reference_group.txt"  # expected by spec; may not exist

    def _parse_reference_label(text: str) -> Optional[str]:
        for line in text.splitlines():
            m = re.match(r"\s*reference\s*=\s*(.+)\s*$", line)
            if m:
                return m.group(1).strip()
        # fallback: first non-empty line
        for line in text.splitlines():
            if line.strip():
                return line.strip()
        return None

    ref_label = None
    if _exists(ref_path):
        ref_label = _parse_reference_label(read_text_any(ref_path))
    else:
        # Strict validation: do NOT auto-create. Missing file must be reported.
        specs.append(
            TableSpec(
                caption=f"Multi-group structural model by {w_label}: reference group metadata.",
                dataframe=None,
                missing_paths=[ref_path],
            )
        )

    missing: List[Path] = []
    if not _exists(pe_path):
        missing.append(pe_path)
        specs.append(
            TableSpec(
                caption=f"Multi-group structural model by {w_label}: key paths by group.",
                dataframe=None,
                missing_paths=missing,
            )
        )
        return specs

    df = cast(pd.DataFrame, read_table_any(pe_path))
    df = normalize_columns(df)

    # Key structural labels tend to end with these tokens
    key_tokens = {"a1", "a2", "b1", "b2", "c", "cc", "cz", "a1c", "a1z", "a2c", "a2z"}

    d = df.copy()
    if not {"op", "lhs", "rhs"}.issubset(d.columns):
        return specs

    # group-specific structural paths
    paths = d[d["op"].astype(str) == "~"].copy()
    if "label" in paths.columns:
        paths["label"] = paths["label"].astype(str)
        paths = paths[paths["label"].str.strip() != ""].copy()
        paths["_tok"] = paths["label"].str.extract(r"([A-Za-z0-9]+)\s*$", expand=False)
        paths = paths[paths["_tok"].isin(key_tokens)].copy()

    paths["Path"] = paths["lhs"].astype(str) + " ~ " + paths["rhs"].astype(str)

    out_cols = [c for c in ["group", "Path", "label", "est", "se", "pvalue", "ci.lower", "ci.upper"] if c in paths.columns]
    paths_out = paths.loc[:, out_cols].copy() if out_cols else pd.DataFrame()
    rename = {"group": "Group", "label": "Label", "est": "Estimate", "se": "SE", "pvalue": "p"}
    paths_out = paths_out.rename(columns=rename)
    if "Estimate" in paths_out.columns:
        paths_out["Estimate"] = as_numeric(paths_out["Estimate"]).map(lambda v: fmt_num(v, nd=3) if pd.notna(v) else "")
    if "SE" in paths_out.columns:
        paths_out["SE"] = as_numeric(paths_out["SE"]).map(lambda v: fmt_num(v, nd=3) if pd.notna(v) else "")
    if "p" in paths_out.columns:
        paths_out["p"] = as_numeric(paths_out["p"]).map(fmt_p)
    if "ci.lower" in paths_out.columns and "ci.upper" in paths_out.columns:
        paths_out["95% CI"] = [fmt_ci(a, b, nd=3) for a, b in zip(paths_out["ci.lower"], paths_out["ci.upper"])]
        paths_out.drop(columns=["ci.lower", "ci.upper"], inplace=True)

    cap = f"Multi-group structural model by {w_label}: key structural paths by group."
    if ref_label:
        cap += f" Reference group: {ref_label}."

    specs.append(TableSpec(caption=cap, dataframe=paths_out if not paths_out.empty else None, missing_paths=[]))
    if paths_out.empty:
        specs[-1].missing_paths = [pe_path]

    # contrasts/defined parameters
    defs = _extract_defined(df)
    if not defs.empty:
        # try to focus on contrast-like definitions
        mask = defs["Defined parameter"].astype(str).str.contains(r"diff|contrast", case=False, regex=True)
        defs_focus = cast(pd.DataFrame, defs.loc[mask, :].copy())
        if defs_focus.empty:
            defs_focus = cast(pd.DataFrame, defs)
        specs.append(
            TableSpec(
                caption=f"Multi-group structural model by {w_label}: defined parameters (including contrasts where available).",
                dataframe=defs_focus,
                missing_paths=[],
            )
        )
    else:
        specs.append(
            TableSpec(
                caption=f"Multi-group structural model by {w_label}: defined parameters (including contrasts where available).",
                dataframe=None,
                missing_paths=[pe_path],
            )
        )

    return specs


def add_missing_output_page(doc: DocxDocument, table_num: int, missing_paths: Sequence[Path]) -> None:
    cap = f"Table {table_num}. Missing Output"
    p = doc.add_paragraph(cap)
    p.runs[0].bold = True

    df = pd.DataFrame({"Missing file": [str(p) for p in missing_paths]})
    add_dataframe_table(doc, df)



def set_doc_default_style(doc: DocxDocument) -> None:
    # python-docx typing stubs are incomplete; cast to Any for style/font access.
    style = cast(Any, doc.styles["Normal"])
    style.font.name = "Times New Roman"
    # Ensure rPr exists
    if getattr(style, "_element", None) is not None and getattr(style._element, "rPr", None) is not None:
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Times New Roman")
    style.font.size = Pt(12)


def add_dataframe_table(doc: DocxDocument, df: pd.DataFrame) -> None:
    df = df.fillna("")
    cols = list(df.columns)

    table = doc.add_table(rows=1, cols=len(cols))
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"

    # header
    hdr = table.rows[0].cells
    for j, c in enumerate(cols):
        run = hdr[j].paragraphs[0].add_run(str(c))
        run.bold = True
        hdr[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER

    # rows
    numeric_cols = set()
    for c in cols:
        # heuristically treat typical numeric columns as numeric for alignment
        if re.search(r"(Estimate|SE|Std|CI|Value|df|chisq|cfi|tli|rmsea|srmr|z|p)$", str(c)):
            numeric_cols.add(c)

    for _, row in df.iterrows():
        cells = table.add_row().cells
        for j, c in enumerate(cols):
            txt = str(row[c]) if row[c] is not None else ""
            para = cells[j].paragraphs[0]
            para.text = txt
            if c in numeric_cols:
                para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
            else:
                para.alignment = WD_ALIGN_PARAGRAPH.LEFT

    # make columns readable
    for col in table.columns:
        for cell in col.cells:
            for par in cell.paragraphs:
                for run in par.runs:
                    run.font.size = Pt(11)


def add_table_page(
    doc: DocxDocument,
    table_num: int,
    caption: str,
    df: Optional[pd.DataFrame],
    missing: Sequence[Path],
    *,
    page_breaks: bool,
) -> int:
    """Add either the table page, or a Missing Output page if needed."""
    if missing:
        add_missing_output_page(doc, table_num, missing)
        if page_breaks:
            doc.add_page_break()
        else:
            doc.add_paragraph("")
        return table_num + 1

    cap = f"Table {table_num}. {caption}"
    p = doc.add_paragraph(cap)
    p.runs[0].bold = True

    if df is None or df.empty:
        # Strict validation: empty tables are treated as missing output.
        add_missing_output_page(doc, table_num, [Path("UNKNOWN_EMPTY_TABLE_SOURCE.txt")])
        if page_breaks:
            doc.add_page_break()
        else:
            doc.add_paragraph("")
        return table_num + 1

    add_dataframe_table(doc, df)

    if page_breaks:
        doc.add_page_break()
    else:
        doc.add_paragraph("")
    return table_num + 1


def main() -> None:
    parser = argparse.ArgumentParser(description="Build Paper_Tables_All.docx from a completed run output folder")
    parser.add_argument(
        "--base_dir",
        type=str,
        default=str(Path(str(Path.cwd())) / DEFAULT_BASE_DIR) if not DEFAULT_BASE_DIR.is_absolute() else str(DEFAULT_BASE_DIR),
        help="Base output directory (e.g., results/.../official_all_RQs_...)",
    )
    parser.add_argument(
        "--out",
        type=str,
        default="",
        help="Optional output docx path (defaults to <base_dir>/Paper_Tables_All.docx)",
    )
    parser.add_argument(
        "--page_breaks",
        type=int,
        default=1,
        help="1 = force one table per page (default). 0 = flow tables continuously to pack pages.",
    )
    args = parser.parse_args()

    base_dir = Path(args.base_dir).expanduser()
    out_docx = Path(args.out).expanduser() if args.out else (base_dir / "Paper_Tables_All.docx")

    if not base_dir.exists():
        raise SystemExit(f"Base output directory not found: {base_dir}")

    doc = DocumentFactory()
    set_doc_default_style(doc)

    page_breaks = bool(int(args.page_breaks))

    table_num = 1

    # Canonical paths for the completed run
    run_log = base_dir / "run_log.txt"
    psw_stage = base_dir / "RQ1_RQ3_main" / "psw_stage_report.txt"
    balance = base_dir / "RQ1_RQ3_main" / "psw_balance_smd.txt"
    rep_data = base_dir / "RQ1_RQ3_main" / "rep_data_with_psw.csv"

    treatment_var = infer_treatment_var(run_log, rep_data)

    fit_main = base_dir / "RQ1_RQ3_main" / "structural" / "structural_fitMeasures.txt"
    r2_main = base_dir / "RQ1_RQ3_main" / "structural" / "structural_r2.txt"
    pe_main = base_dir / "RQ1_RQ3_main" / "structural" / "structural_parameterEstimates.txt"
    stdsol_main = base_dir / "RQ1_RQ3_main" / "structural" / "structural_standardizedSolution.txt"

    # Sensitivity outputs (expected to be created by a lightweight unweighted fit)
    pe_unweighted = base_dir / "sensitivity_unweighted_parallel" / "structural" / "structural_parameterEstimates.txt"

    meas_root = base_dir / "RQ4_measurement"
    race_root = base_dir / "RQ4_structural_by_re_all"

    # Extract W list from run_log if present
    w_vars_meas: List[str] = []
    if _exists(run_log):
        txt = read_text_any(run_log)
        m = re.search(r"W_VARS_MEAS_OK:\s*(.+)", txt)
        if m:
            w_vars_meas = [t.strip() for t in m.group(1).split(",") if t.strip()]

    # Extract PS model covariates for Table 2 note
    ps_model_note = ""
    if _exists(psw_stage):
        ps_line = _parse_ps_model_line(read_text_any(psw_stage))
        if ps_line:
            covs = _parse_ps_covariates(ps_line)
            if covs:
                ps_model_note = "PS model covariates: " + ", ".join(covs) + "."

    # Table set (EXACTLY 12):
    key_labels = ["a1", "a2", "b1", "b2", "c", "a1c", "a1z", "a2c", "a2z", "cc", "cz"]

    specs: List[TableSpec] = []

    # Table 1
    specs.append(table_sample_split(rep_data, treatment_var=treatment_var))

    # Table 2
    bal = table_balance(balance)
    if ps_model_note:
        bal.caption = bal.caption + " Note. " + ps_model_note
    specs.append(bal)

    # Table 3
    specs.append(compute_weight_diagnostics(psw_stage, rep_data))

    # Table 4A
    specs.append(table_fit_indices(fit_main, caption="Global fit indices for the primary pooled model (MLR; robust/scaled reported where available)."))

    # Table 4B
    specs.append(table_standardized_loadings(stdsol_main))

    # Table 5
    specs.append(table_invariance_summary(meas_root, w_vars=w_vars_meas if w_vars_meas else ["re_all", "firstgen", "pell", "sex", "living18"]))

    # Table 6
    specs.append(table_r2_summary(r2_main))

    # Table 7
    t7 = table_structural_paths(
        pe_main,
        caption="Primary pooled structural paths with bias-corrected bootstrap CIs (bca.simple).",
        include_labels=key_labels,
    )
    if t7.dataframe is not None and "95% CI" in t7.dataframe.columns:
        t7.dataframe = t7.dataframe.rename(columns={"95% CI": "Bias-corrected bootstrap CI (bca.simple)"})
    specs.append(t7)

    # Table 8 (conditional effects at Z)
    # Build from defined params
    dspec = table_defined_params(pe_main, caption="")
    if dspec.dataframe is not None and not dspec.dataframe.empty:
        dp = dspec.dataframe.copy()
        names = dp["Defined parameter"].astype(str)
        want = {
            "dir_z_low": ("Direct", "Low"),
            "dir_z_mid": ("Direct", "Mid"),
            "dir_z_high": ("Direct", "High"),
            "ind_EmoDiss_z_low": ("Indirect (via EmoDiss)", "Low"),
            "ind_EmoDiss_z_mid": ("Indirect (via EmoDiss)", "Mid"),
            "ind_EmoDiss_z_high": ("Indirect (via EmoDiss)", "High"),
            "ind_QualEngag_z_low": ("Indirect (via QualEngag)", "Low"),
            "ind_QualEngag_z_mid": ("Indirect (via QualEngag)", "Mid"),
            "ind_QualEngag_z_high": ("Indirect (via QualEngag)", "High"),
            "total_z_low": ("Total", "Low"),
            "total_z_mid": ("Total", "Mid"),
            "total_z_high": ("Total", "High"),
        }
        keep_rows = dp[names.isin(set(want.keys()))].copy()
        if not keep_rows.empty:
            keep_rows["Effect"] = keep_rows["Defined parameter"].map(lambda s: want.get(str(s), ("", ""))[0])
            keep_rows["Z"] = keep_rows["Defined parameter"].map(lambda s: want.get(str(s), ("", ""))[1])
            out_cols = [c for c in ["Effect", "Z", "Estimate", "SE", "p", "95% CI"] if c in keep_rows.columns]
            t8_df = keep_rows[out_cols].copy()
        else:
            t8_df = None
    else:
        t8_df = None
    specs.append(
        TableSpec(
            caption="Conditional effects at Z (credit_dose_c): direct, indirect, and total effects at low/mid/high Z.",
            dataframe=t8_df,
            missing_paths=[] if (t8_df is not None and not t8_df.empty) else [pe_main],
        )
    )

    # Table 9 (indices of moderated mediation)
    if dspec.dataframe is not None and not dspec.dataframe.empty:
        dp = dspec.dataframe.copy()
        names = dp["Defined parameter"].astype(str)
        keep_rows = dp[names.str.startswith("index_MM_")].copy()
        t9_df = keep_rows[[c for c in ["Defined parameter", "Estimate", "SE", "p", "95% CI"] if c in keep_rows.columns]].copy() if not keep_rows.empty else None
    else:
        t9_df = None
    specs.append(
        TableSpec(
            caption="Index of moderated mediation (IMM terms).",
            dataframe=t9_df,
            missing_paths=[] if (t9_df is not None and not t9_df.empty) else [pe_main],
        )
    )

    # Table 10 (race fit)
    specs.append(table_race_fit_indices(race_root))

    # Table 10B (race key coefficients)
    specs.append(table_race_key_paths(race_root, key_labels=key_labels))

    # Table 11 (descriptives)
    specs.append(table_descriptives_by_treatment(rep_data, treatment_var=treatment_var))

    # Table 12 (sensitivity)
    specs.append(table_sensitivity_unweighted_vs_weighted(pe_main, pe_unweighted, key_labels=key_labels))

    # ---- Emit pages
    for spec in specs:
        table_num = add_table_page(
            doc,
            table_num,
            spec.caption,
            spec.dataframe,
            spec.missing_paths,
            page_breaks=page_breaks,
        )

    out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_docx))
    print(f"Wrote: {out_docx}")

    # Audit trail: list exactly what inputs were read.
    used = sorted({str(p) for p in USED_INPUT_PATHS})
    print("Inputs used:")
    for p in used:
        print(f"- {p}")


if __name__ == "__main__":
    main()
