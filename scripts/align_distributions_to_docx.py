#!/usr/bin/env python3
"""Align marginal distributions in a CSV to targets in a DOCX table.

This is intended for *synthetic / simulated* datasets where you want the
marginal distributions (e.g., race %, sex %, grade buckets, etc.) to match
a reference table.

It makes a minimally-invasive adjustment: it only changes as many rows as
necessary per variable to match the target counts (within rounding).

Outputs:
- adjusted dataset CSV
- summary CSV with before/after proportions
- notes TXT with parsing assumptions

Usage:
  python scripts/align_distributions_to_docx.py --in rep_data.csv --table "dataset distribu.docx" \
    --out results/distribution_alignment/rep_data_aligned.csv --seed 123
"""

from __future__ import annotations

import argparse
import re
from dataclasses import dataclass
from pathlib import Path
from typing import Dict, Iterable, List, Optional, Sequence, Tuple, cast

import numpy as np
import pandas as pd
from docx import Document


def _parse_percent(text: str) -> Optional[float]:
    """Parse a percent-like cell into a numeric percent (0-100).

    Handles examples like:
    - "47.7 %"
    - "≈ 51.5 % (A/A+) + 27 % (A–)"
    - "20–30 %"
    - "0.5 % total"

    Returns None if no numbers found.
    """

    s = (text or "").strip()
    if not s:
        return None

    s = s.replace(",", "")

    # Range like 20–30 % or 20-30 %
    m = re.search(r"(\d+(?:\.\d+)?)\s*[\-–]\s*(\d+(?:\.\d+)?)", s)
    if m:
        lo = float(m.group(1))
        hi = float(m.group(2))
        return (lo + hi) / 2.0

    # Sum of multiple percents
    nums = re.findall(r"\d+(?:\.\d+)?", s)
    if not nums:
        return None

    return float(sum(float(x) for x in nums))


@dataclass(frozen=True)
class TargetSpec:
    column: str
    # map from dataset level -> target proportion (0..1)
    level_probs: Dict[str, float]
    notes: str


def _normalize_probs(level_probs: Dict[str, float]) -> Dict[str, float]:
    total = float(sum(level_probs.values()))
    if total <= 0:
        raise ValueError("Target probabilities sum to 0")
    return {k: float(v) / total for k, v in level_probs.items()}


def _adjust_categorical_min_change(
    s: pd.Series,
    target_probs: Dict[str, float],
    rng: np.random.Generator,
) -> Tuple[pd.Series, Dict[str, int]]:
    """Adjust a categorical series to match target_probs with minimal changes.

    Only levels that already exist in either the target or series are considered.
    Levels not mentioned in target_probs keep their relative frequencies within the
    leftover mass *only if* you pass them explicitly; otherwise they are treated as 0.
    """

    # Work on a copy
    out = s.copy()

    # Determine universe of levels we will control
    current_counts = out.value_counts(dropna=False)
    levels = list(target_probs.keys())

    n = int(len(out))
    desired_counts: Dict[str, int] = {}

    # Convert probs -> integer counts; remainder assigned deterministically
    raw = {lvl: float(target_probs[lvl]) * n for lvl in levels}
    floor_counts = {lvl: int(np.floor(raw[lvl])) for lvl in levels}
    remainder = n - int(sum(floor_counts.values()))
    # Distribute remainder to largest fractional parts
    fracs = sorted(((raw[lvl] - floor_counts[lvl], lvl) for lvl in levels), reverse=True)
    desired_counts = dict(floor_counts)
    for i in range(remainder):
        desired_counts[fracs[i % len(fracs)][1]] += 1

    # Build surplus pool and deficit list
    surplus_indices: List[int] = []
    deficit_plan: List[Tuple[str, int]] = []

    for lvl in levels:
        cur = int(current_counts.get(lvl, 0))
        des = int(desired_counts[lvl])
        if cur > des:
            idx = out.index[out == lvl].to_numpy()
            # choose which rows to change (random for now)
            choose = rng.choice(idx, size=(cur - des), replace=False)
            surplus_indices.extend(int(i) for i in choose)
        elif cur < des:
            deficit_plan.append((lvl, des - cur))

    # If target includes levels not present, we can create them by reassigning from surplus.
    # If there is no surplus (shouldn't happen if targets sum to 1), we will reassign anyway.
    rng.shuffle(surplus_indices)

    cursor = 0
    changes = 0
    for lvl, need in deficit_plan:
        take = min(need, len(surplus_indices) - cursor)
        if take < need:
            # Not enough surplus (e.g., target uses unseen levels but rounding issues);
            # fall back to taking from any other level.
            remaining = need - take
            pool = out.index.to_numpy()
            extra = rng.choice(pool, size=remaining, replace=False)
            for i in extra:
                if out.at[i] != lvl:
                    out.at[i] = lvl
                    changes += 1
        for j in range(cursor, cursor + take):
            i = surplus_indices[j]
            if out.at[i] != lvl:
                out.at[i] = lvl
                changes += 1
        cursor += take

    change_counts = {"rows_changed": int(changes)}
    return out, change_counts


def _adjust_binary(
    s: pd.Series,
    target_p1: float,
    rng: np.random.Generator,
    one_value: int = 1,
    zero_value: int = 0,
) -> Tuple[pd.Series, Dict[str, int]]:
    out = s.copy()
    n = len(out)
    desired_ones = int(round(float(target_p1) * n))
    cur_ones = int((out == one_value).sum())
    if cur_ones == desired_ones:
        return out, {"rows_changed": 0}

    if desired_ones > cur_ones:
        idx0 = out.index[out == zero_value].to_numpy()
        need = desired_ones - cur_ones
        choose = rng.choice(idx0, size=need, replace=False)
        out.loc[choose] = one_value
        return out, {"rows_changed": int(need)}

    idx1 = out.index[out == one_value].to_numpy()
    need = cur_ones - desired_ones
    choose = rng.choice(idx1, size=need, replace=False)
    out.loc[choose] = zero_value
    return out, {"rows_changed": int(need)}


def _adjust_likert_agree(
    s: pd.Series,
    target_agree: float,
    agree_threshold: int,
    rng: np.random.Generator,
) -> Tuple[pd.Series, Dict[str, int]]:
    """Adjust Likert-type responses to hit target proportion >= agree_threshold."""
    out = s.copy()
    n = len(out)
    desired_agree = int(round(float(target_agree) * n))
    is_agree = out >= agree_threshold
    cur_agree = int(is_agree.sum())
    if cur_agree == desired_agree:
        return out, {"rows_changed": 0}

    if desired_agree > cur_agree:
        need = desired_agree - cur_agree
        pool = out.index[~is_agree].to_numpy()
        choose = rng.choice(pool, size=need, replace=False)
        # Promote to the threshold (minimal shift)
        out.loc[choose] = agree_threshold
        return out, {"rows_changed": int(need)}

    need = cur_agree - desired_agree
    pool = out.index[is_agree].to_numpy()
    choose = rng.choice(pool, size=need, replace=False)
    # Demote to just below agree threshold (minimal shift)
    out.loc[choose] = agree_threshold - 1
    return out, {"rows_changed": int(need)}


def _read_docx_targets(docx_path: Path) -> Tuple[pd.DataFrame, List[str]]:
    doc = Document(str(docx_path))
    if not doc.tables:
        raise ValueError(f"No tables found in {docx_path}")

    t = doc.tables[0]
    rows: List[List[str]] = []
    for r in t.rows:
        rows.append([c.text.strip().replace("\n", " ") for c in r.cells])

    header = rows[0]
    data = rows[1:]

    df = pd.DataFrame(data, columns=header)

    # forward-fill category / variable name
    cat_col = header[0]
    df[cat_col] = df[cat_col].replace("", np.nan).ffill()

    notes: List[str] = []
    notes.append(f"DOCX: used first table only; columns={header}")
    return df, notes


def _build_specs_from_table(table: pd.DataFrame) -> Tuple[List[TargetSpec], List[str]]:
    cat = table.columns[0]
    level = table.columns[1]
    pct = table.columns[2]

    notes: List[str] = []

    # Helper to collect rows for a category
    def rows_for(category: str) -> pd.DataFrame:
        return table.loc[table[cat].str.strip() == category, [level, pct]].copy()

    specs: List[TargetSpec] = []

    # High School Grades (A–F) -> hgrades_AF
    hs = rows_for("High School Grades (A–F)")
    if not hs.empty:
        p_a = _parse_percent(str(hs.loc[hs[level].str.contains("Mostly A", na=False), pct].iloc[0]))
        p_b = _parse_percent(str(hs.loc[hs[level].str.contains("Mostly B", na=False), pct].iloc[0]))
        if p_a is not None and p_b is not None:
            # remainder goes to C/D/F (distributed proportional to observed in data later)
            probs = {"A": p_a / 100.0, "B": p_b / 100.0}
            specs.append(TargetSpec(column="hgrades_AF", level_probs=probs, notes="Remainder allocated across C/D/F proportional to current C/D/F."))
        else:
            notes.append("High School Grades: could not parse A/B percents")
    else:
        notes.append("High School Grades: category not found")

    # Race / Ethnicity -> re_all (collapse multiple small groups into Other/Multiracial/Unknown)
    race = rows_for("Race / Ethnicity")
    if not race.empty:
        def pct_for(label: str) -> Optional[float]:
            m = race.loc[race[level].str.strip() == label, pct]
            if m.empty:
                return None
            return _parse_percent(str(m.iloc[0]))

        p_hisp = pct_for("Hispanic / Latino")
        p_white = pct_for("White")
        p_asian = pct_for("Asian")
        p_two = pct_for("Two or More Races")
        p_black = pct_for("Black / African American")
        p_unknown = pct_for("Race Unknown")
        p_intl = pct_for("International")
        p_small = pct_for("NH/PI + AI/AN")

        if None not in (p_hisp, p_white, p_asian, p_black):
            # Help type-checkers: values are guaranteed not-None in this branch.
            p_hisp_f = cast(float, p_hisp)
            p_white_f = cast(float, p_white)
            p_asian_f = cast(float, p_asian)
            p_black_f = cast(float, p_black)
            other_sum = sum(x for x in [p_two, p_unknown, p_intl, p_small] if x is not None)
            probs = {
                "Hispanic/Latino": p_hisp_f / 100.0,
                "White": p_white_f / 100.0,
                "Asian": p_asian_f / 100.0,
                "Black/African American": p_black_f / 100.0,
                "Other/Multiracial/Unknown": float(other_sum) / 100.0,
            }
            probs = _normalize_probs(probs)
            specs.append(TargetSpec(column="re_all", level_probs=probs, notes="Collapsed Two+ / Unknown / International / NH/PI+AI/AN into Other/Multiracial/Unknown; normalized to sum to 1."))
        else:
            notes.append("Race: missing required race percents (Hispanic/White/Asian/Black)")
    else:
        notes.append("Race: category not found")

    # Living situation -> living18
    living = rows_for("Living Situation (FY students)")
    if not living.empty:
        def pct_for_contains(substr: str) -> Optional[float]:
            m = living.loc[living[level].str.contains(substr, na=False), pct]
            if m.empty:
                return None
            return _parse_percent(str(m.iloc[0]))

        p_on = pct_for_contains("On-campus")
        p_off = pct_for_contains("Off-campus")
        p_fam = pct_for_contains("family")
        if None not in (p_on, p_off, p_fam):
            # Help type-checkers: values are guaranteed not-None in this branch.
            p_on_f = cast(float, p_on)
            p_off_f = cast(float, p_off)
            p_fam_f = cast(float, p_fam)
            probs = {
                "On-campus (residence hall)": p_on_f / 100.0,
                "Off-campus (rent/apartment)": p_off_f / 100.0,
                "With family (commuting)": p_fam_f / 100.0,
            }
            probs = _normalize_probs(probs)
            specs.append(TargetSpec(column="living18", level_probs=probs, notes="Used midpoints for ranges; normalized to sum to 1."))
        else:
            notes.append("Living: could not parse all 3 percentages")

    # Sex / Gender -> sex
    sex = rows_for("Sex / Gender")
    if not sex.empty:
        p_f = _parse_percent(str(sex.loc[sex[level].str.strip() == "Female", pct].iloc[0]))
        p_m = _parse_percent(str(sex.loc[sex[level].str.strip() == "Male", pct].iloc[0]))
        if p_f is not None and p_m is not None:
            probs = {"Woman": p_f / 100.0, "Man": p_m / 100.0}
            probs = _normalize_probs(probs)
            specs.append(TargetSpec(column="sex", level_probs=probs, notes="Mapped Female->Woman, Male->Man; normalized to sum to 1."))

    # Cohort indicator -> cohort (binary)
    cohort = rows_for("Cohort Indicator (EOP etc.)")
    if not cohort.empty:
        p_eop = _parse_percent(str(cohort.loc[cohort[level].str.contains("EOP", na=False), pct].iloc[0]))
        if p_eop is not None:
            specs.append(TargetSpec(column="cohort", level_probs={"1": p_eop / 100.0, "0": 1.0 - (p_eop / 100.0)}, notes="Assumed cohort==1 corresponds to EOP/Bridge participants."))

    # Pell -> pell (binary)
    pell = rows_for("Pell Grant Status")
    if not pell.empty:
        p_rec = _parse_percent(str(pell.loc[pell[level].str.contains("Recipients", na=False), pct].iloc[0]))
        if p_rec is not None:
            specs.append(TargetSpec(column="pell", level_probs={"1": p_rec / 100.0, "0": 1.0 - (p_rec / 100.0)}, notes="Assumed pell==1 is recipient."))

    # AP Courses -> hapcl (binary)
    ap = rows_for("AP Courses Completed")
    if not ap.empty:
        p_gt2 = _parse_percent(str(ap.loc[ap[level].str.contains(">", na=False), pct].iloc[0]))
        if p_gt2 is not None:
            specs.append(TargetSpec(column="hapcl", level_probs={"1": p_gt2 / 100.0, "0": 1.0 - (p_gt2 / 100.0)}, notes="Assumed hapcl==1 corresponds to >2 AP courses."))

    # First-gen -> firstgen (binary)
    fg = rows_for("First-Generation Status")
    if not fg.empty:
        p_fg = _parse_percent(str(fg.loc[fg[level].str.contains("First-generation", na=False), pct].iloc[0]))
        if p_fg is not None:
            specs.append(TargetSpec(column="firstgen", level_probs={"1": p_fg / 100.0, "0": 1.0 - (p_fg / 100.0)}, notes="Assumed firstgen==1 corresponds to first-generation."))

    # SB items -> sbmyself, sbvalued (agree/strongly agree)
    sb1 = rows_for("Sense of Belonging – “I can be myself”")
    if not sb1.empty:
        p_agree = _parse_percent(str(sb1.loc[sb1[level].str.contains("Agree", na=False), pct].iloc[0]))
        if p_agree is not None:
            specs.append(TargetSpec(column="sbmyself", level_probs={"agree": p_agree / 100.0}, notes="Interpreted Agree/Strongly Agree as >=4 on a 1-5 scale."))

    sb2 = rows_for("Sense of Belonging – “I feel valued”")
    if not sb2.empty:
        p_agree = _parse_percent(str(sb2.loc[sb2[level].str.contains("Agree", na=False), pct].iloc[0]))
        if p_agree is not None:
            specs.append(TargetSpec(column="sbvalued", level_probs={"agree": p_agree / 100.0}, notes="Interpreted Agree/Strongly Agree as >=4 on a 1-5 scale."))

    return specs, notes


def main() -> int:
    ap = argparse.ArgumentParser()
    ap.add_argument("--in", dest="in_csv", required=True, help="Input CSV (e.g., rep_data.csv)")
    ap.add_argument("--table", dest="docx", required=True, help="DOCX file containing distribution table")
    ap.add_argument("--out", dest="out_csv", required=True, help="Output adjusted CSV")
    ap.add_argument("--summary", default=None, help="Output summary CSV (before/after)")
    ap.add_argument("--notes", default=None, help="Output notes TXT")
    ap.add_argument("--seed", type=int, default=123, help="RNG seed")
    args = ap.parse_args()

    in_csv = Path(args.in_csv)
    docx = Path(args.docx)
    out_csv = Path(args.out_csv)
    out_csv.parent.mkdir(parents=True, exist_ok=True)

    if not in_csv.exists():
        raise FileNotFoundError(f"Input CSV not found: {in_csv}")
    if not docx.exists():
        raise FileNotFoundError(
            f"DOCX table file not found: {docx}. "
            "Pass the correct path via --table (the file must exist on disk)."
        )

    summary_csv = Path(args.summary) if args.summary else out_csv.with_name("alignment_summary.csv")
    notes_txt = Path(args.notes) if args.notes else out_csv.with_name("alignment_notes.txt")

    rng = np.random.default_rng(args.seed)

    df = pd.read_csv(in_csv)
    table, doc_notes = _read_docx_targets(docx)
    specs, spec_notes = _build_specs_from_table(table)

    notes: List[str] = []
    notes.extend(doc_notes)
    notes.extend(spec_notes)

    summary_rows: List[Dict[str, object]] = []

    def record_before_after(col: str, before: pd.Series, after: pd.Series, target_desc: str) -> None:
        if pd.api.types.is_numeric_dtype(before):
            summary_rows.append({
                "column": col,
                "metric": "mean",
                "before": float(before.mean()),
                "after": float(after.mean()),
                "target": target_desc,
            })
        # Always record categorical proportions for top levels
        bvc = before.astype(str).value_counts(normalize=True)
        avc = after.astype(str).value_counts(normalize=True)
        levels = sorted(set(bvc.index).union(set(avc.index)))
        for lvl in levels:
            summary_rows.append({
                "column": col,
                "metric": f"prop[{lvl}]",
                "before": float(bvc.get(lvl, 0.0)),
                "after": float(avc.get(lvl, 0.0)),
                "target": target_desc,
            })

    for spec in specs:
        col = spec.column
        if col not in df.columns:
            notes.append(f"SKIP: column '{col}' not found in dataset")
            continue

        before = df[col]

        # Special-case: Likert agree targets
        if col in ("sbmyself", "sbvalued") and "agree" in spec.level_probs:
            after, change = _adjust_likert_agree(before.astype(int), float(spec.level_probs["agree"]), agree_threshold=4, rng=rng)
            df[col] = after
            notes.append(f"ADJUST: {col} agree>=4 to {spec.level_probs['agree']:.3f}; {change['rows_changed']} rows changed")
            record_before_after(col, before, after, f"agree>=4={spec.level_probs['agree']:.3f}")
            continue

        # Binary numeric columns
        if pd.api.types.is_numeric_dtype(before) and set(spec.level_probs.keys()) <= {"0", "1"}:
            target_p1 = float(spec.level_probs.get("1", 0.0))
            after, change = _adjust_binary(before.astype(int), target_p1=target_p1, rng=rng)
            df[col] = after
            notes.append(f"ADJUST: {col} P(1)={target_p1:.3f}; {change['rows_changed']} rows changed")
            record_before_after(col, before, after, f"P(1)={target_p1:.3f}")
            continue

        # Categorical (strings)
        # For hgrades_AF we allow remainder allocation across C/D/F proportional to current
        if col == "hgrades_AF":
            probs = dict(spec.level_probs)
            # Allocate remainder among C/D/F proportional to current within those levels
            remainder = 1.0 - float(sum(probs.values()))
            if remainder < 0:
                probs = _normalize_probs(probs)
                remainder = 0.0
            cur_cdf = before.value_counts(normalize=True)
            tail_levels = [lvl for lvl in ["C", "D", "F"] if lvl in cur_cdf.index]
            # Compute explicitly to avoid type-checker ambiguity around Series indexing.
            tail_mass = float(sum(float(cur_cdf[lvl]) for lvl in tail_levels)) if tail_levels else 0.0
            if tail_levels and remainder > 0 and tail_mass > 0:
                for lvl in tail_levels:
                    probs[lvl] = remainder * float(cur_cdf[lvl] / tail_mass)
            else:
                # fallback: put remainder into C if present
                if remainder > 0:
                    probs["C"] = probs.get("C", 0.0) + remainder
            probs = _normalize_probs(probs)
            after, change = _adjust_categorical_min_change(before.astype(str), probs, rng)
            df[col] = after
            notes.append(f"ADJUST: {col} targets={probs}; {change['rows_changed']} rows changed")
            record_before_after(col, before.astype(str), after.astype(str), f"targets={probs}")
            continue

        # Standard categorical target
        probs = _normalize_probs(dict(spec.level_probs))
        after, change = _adjust_categorical_min_change(before.astype(str), probs, rng)
        df[col] = after
        notes.append(f"ADJUST: {col} targets={probs}; {change['rows_changed']} rows changed")
        record_before_after(col, before.astype(str), after.astype(str), f"targets={probs}")

    df.to_csv(out_csv, index=False)
    pd.DataFrame(summary_rows).to_csv(summary_csv, index=False)

    notes.append("")
    notes.append("UNMAPPED TABLE ROWS (no dataset column found):")
    notes.append("- High School Type (no obvious column in rep_data.csv)")

    notes_txt.write_text("\n".join(notes), encoding="utf-8")

    print(f"Wrote adjusted CSV: {out_csv}")
    print(f"Wrote summary: {summary_csv}")
    print(f"Wrote notes: {notes_txt}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
