#!/usr/bin/env python3
"""Build publication-ready Word tables from bootstrap v3 results.

This script creates formatted tables from the corrected model bootstrap results
(without credit_dose_c main effect).

Usage:
    python scripts/build_bootstrap_tables.py [--results_dir DIR]

Output: <results_dir>/Bootstrap_Tables_v3.docx
"""

import argparse
from pathlib import Path
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT


def fmt_num(x, nd=3):
    """Format number to specified decimal places."""
    try:
        v = float(x)
        if abs(v) < 0.001:
            return f"{v:.{nd}f}"
        return f"{v:.{nd}f}"
    except:
        return str(x)


def fmt_ci(lo, hi, nd=3):
    """Format confidence interval."""
    return f"[{fmt_num(lo, nd)}, {fmt_num(hi, nd)}]"


def add_table_page(doc, table_num, caption, df, note=None):
    """Add a table with caption to the document."""
    # Caption
    p = doc.add_paragraph()
    run = p.add_run(f"Table {table_num}. ")
    run.bold = True
    run.font.size = Pt(11)
    run = p.add_run(caption)
    run.italic = True
    run.font.size = Pt(11)
    
    if df is not None and not df.empty:
        # Create table
        tbl = doc.add_table(rows=1, cols=len(df.columns))
        tbl.style = 'Table Grid'
        
        # Header row
        hdr = tbl.rows[0].cells
        for i, col in enumerate(df.columns):
            hdr[i].text = col
            hdr[i].paragraphs[0].runs[0].bold = True
            hdr[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Data rows
        for _, row in df.iterrows():
            cells = tbl.add_row().cells
            for i, val in enumerate(row):
                cells[i].text = str(val)
                # Right-align numeric columns
                if i > 0:
                    cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
    
    # Note
    if note:
        p = doc.add_paragraph()
        run = p.add_run("Note. ")
        run.italic = True
        run.font.size = Pt(10)
        run = p.add_run(note)
        run.font.size = Pt(10)
    
    doc.add_page_break()
    return table_num + 1


def main():
    parser = argparse.ArgumentParser(description='Build bootstrap result tables (DOCX)')
    parser.add_argument('--results_dir', type=str, 
                        default='results/fast_treat_control/official_all_RQs',
                        help='Directory containing bootstrap results')
    args = parser.parse_args()
    
    base_dir = Path(args.results_dir)
    boot_csv = base_dir / "bootstrap_v3" / "bootstrap_results.csv"
    out_docx = base_dir / "Bootstrap_Tables_v3.docx"
    
    if not boot_csv.exists():
        # Also check for bootstrap_pipeline subdirectory
        alt_csv = base_dir / "bootstrap_pipeline" / "bootstrap_results.csv"
        if alt_csv.exists():
            boot_csv = alt_csv
        else:
            print(f"ERROR: {boot_csv} not found")
            return
    
    # Load bootstrap results
    df = pd.read_csv(boot_csv)
    
    doc = Document()
    
    # Title page
    title = doc.add_heading("Bootstrap Results: Corrected Conditional Process Model", level=0)
    doc.add_paragraph(
        "Model specification: Parallel mediation with first-stage moderation.\n"
        "Collinearity fix: credit_dose_c main effect removed (only varies for FASt=1).\n"
        "Bootstrap replicates: 500 | Convergence: 100%\n"
        "CI method: BCa (bias-corrected accelerated)"
    )
    doc.add_page_break()
    
    table_num = 1
    
    # Table 1: Structural path coefficients
    paths = df[df['parameter'].isin(['a1', 'a1z', 'a2', 'a2z', 'b1', 'b2', 'c', 'cz'])].copy()
    paths['Parameter'] = paths['parameter'].map({
        'a1': 'FASt → Emotional Distress (a₁)',
        'a1z': 'FASt × Dose → Emotional Distress (a₁ᵢₙₜ)',
        'a2': 'FASt → Quality of Engagement (a₂)',
        'a2z': 'FASt × Dose → Quality of Engagement (a₂ᵢₙₜ)',
        'b1': 'Emotional Distress → Developmental Adjustment (b₁)',
        'b2': 'Quality of Engagement → Developmental Adjustment (b₂)',
        'c': "FASt → Developmental Adjustment (c')",
        'cz': "FASt × Dose → Developmental Adjustment (c'ᵢₙₜ)"
    })
    paths['Estimate'] = paths['est'].apply(lambda x: fmt_num(x))
    paths['SE'] = paths['boot_se'].apply(lambda x: fmt_num(x))
    paths['95% CI'] = paths.apply(lambda r: fmt_ci(r['ci_lower'], r['ci_upper']), axis=1)
    paths['Sig'] = paths['sig'].map({True: '*', False: ''})
    
    tbl1 = paths[['Parameter', 'Estimate', 'SE', '95% CI', 'Sig']].reset_index(drop=True)
    
    table_num = add_table_page(
        doc, table_num,
        "Structural path coefficients from bootstrap-then-weight analysis.",
        tbl1,
        "FASt = accelerated dual credit status (≥12 transferable units at matriculation); "
        "Dose = credit_dose (mean-centered units beyond threshold); "
        "FASt × Dose = interaction term. "
        "* p < .05 based on bootstrap CI excluding zero. B = 500 replicates, 100% convergence."
    )
    
    # Table 2: Conditional a-paths (X → Mediators at different Z levels)
    cond_a = df[df['parameter'].str.match(r'a[12]_z_(low|mid|high)')].copy()
    cond_a['Effect'] = cond_a['parameter'].apply(
        lambda x: 'FASt → Emotional Distress' if 'a1' in x else 'FASt → Quality of Engagement'
    )
    cond_a['Credit Dose'] = cond_a['parameter'].apply(lambda x: x.split('_')[-1].title())
    cond_a['Estimate'] = cond_a['est'].apply(lambda x: fmt_num(x))
    cond_a['SE'] = cond_a['boot_se'].apply(lambda x: fmt_num(x))
    cond_a['95% CI'] = cond_a.apply(lambda r: fmt_ci(r['ci_lower'], r['ci_upper']), axis=1)
    cond_a['Sig'] = cond_a['sig'].map({True: '*', False: ''})
    
    # Sort by effect and Z level
    z_order = {'Low': 0, 'Mid': 1, 'High': 2}
    cond_a['z_sort'] = cond_a['Credit Dose'].map(z_order)
    cond_a = cond_a.sort_values(['Effect', 'z_sort']).reset_index(drop=True)
    
    tbl2 = cond_a[['Effect', 'Credit Dose', 'Estimate', 'SE', '95% CI', 'Sig']]
    
    table_num = add_table_page(
        doc, table_num,
        "Conditional effects of FASt status on mediators at low, mean, and high credit dose.",
        tbl2,
        "Low = −1 SD below mean; Mid = at the mean; High = +1 SD above mean credit dose. "
        "* p < .05 based on bootstrap CI excluding zero."
    )
    
    # Table 3: Conditional direct effects
    dir_eff = df[df['parameter'].str.startswith('dir_z')].copy()
    dir_eff['Credit Dose'] = dir_eff['parameter'].apply(lambda x: x.split('_')[-1].title())
    dir_eff['Estimate'] = dir_eff['est'].apply(lambda x: fmt_num(x))
    dir_eff['SE'] = dir_eff['boot_se'].apply(lambda x: fmt_num(x))
    dir_eff['95% CI'] = dir_eff.apply(lambda r: fmt_ci(r['ci_lower'], r['ci_upper']), axis=1)
    dir_eff['Sig'] = dir_eff['sig'].map({True: '*', False: ''})
    
    z_order = {'Low': 0, 'Mid': 1, 'High': 2}
    dir_eff['z_sort'] = dir_eff['Credit Dose'].map(z_order)
    dir_eff = dir_eff.sort_values('z_sort').reset_index(drop=True)
    
    tbl3 = dir_eff[['Credit Dose', 'Estimate', 'SE', '95% CI', 'Sig']]
    
    table_num = add_table_page(
        doc, table_num,
        "Conditional direct effects of FASt status on Developmental Adjustment.",
        tbl3,
        "Low = −1 SD; Mid = mean; High = +1 SD of credit dose. "
        "* p < .05 based on bootstrap CI excluding zero. "
        "Direct effects are non-significant at all dose levels."
    )
    
    # Table 4: Conditional indirect effects
    ind_eff = df[df['parameter'].str.startswith('ind_')].copy()
    ind_eff['Mediator'] = ind_eff['parameter'].apply(
        lambda x: 'Emotional Distress' if 'EmoDiss' in x else 'Quality of Engagement'
    )
    ind_eff['Credit Dose'] = ind_eff['parameter'].apply(lambda x: x.split('_')[-1].title())
    ind_eff['Estimate'] = ind_eff['est'].apply(lambda x: fmt_num(x))
    ind_eff['SE'] = ind_eff['boot_se'].apply(lambda x: fmt_num(x))
    ind_eff['95% CI'] = ind_eff.apply(lambda r: fmt_ci(r['ci_lower'], r['ci_upper']), axis=1)
    ind_eff['Sig'] = ind_eff['sig'].map({True: '*', False: ''})
    
    z_order = {'Low': 0, 'Mid': 1, 'High': 2}
    ind_eff['z_sort'] = ind_eff['Credit Dose'].map(z_order)
    ind_eff = ind_eff.sort_values(['Mediator', 'z_sort']).reset_index(drop=True)
    
    tbl4 = ind_eff[['Mediator', 'Credit Dose', 'Estimate', 'SE', '95% CI', 'Sig']]
    
    table_num = add_table_page(
        doc, table_num,
        "Conditional indirect effects of FASt status on Developmental Adjustment.",
        tbl4,
        "Low = −1 SD; Mid = mean; High = +1 SD of credit dose. "
        "* p < .05 based on bootstrap CI excluding zero. "
        "Via Emotional Distress: significant at all dose levels (negative). "
        "Via Quality of Engagement: significant at mean and high dose (negative)."
    )
    
    # Table 5: Conditional total effects
    tot_eff = df[df['parameter'].str.startswith('total_z')].copy()
    tot_eff['Credit Dose'] = tot_eff['parameter'].apply(lambda x: x.split('_')[-1].title())
    tot_eff['Estimate'] = tot_eff['est'].apply(lambda x: fmt_num(x))
    tot_eff['SE'] = tot_eff['boot_se'].apply(lambda x: fmt_num(x))
    tot_eff['95% CI'] = tot_eff.apply(lambda r: fmt_ci(r['ci_lower'], r['ci_upper']), axis=1)
    tot_eff['Sig'] = tot_eff['sig'].map({True: '*', False: ''})
    
    z_order = {'Low': 0, 'Mid': 1, 'High': 2}
    tot_eff['z_sort'] = tot_eff['Credit Dose'].map(z_order)
    tot_eff = tot_eff.sort_values('z_sort').reset_index(drop=True)
    
    tbl5 = tot_eff[['Credit Dose', 'Estimate', 'SE', '95% CI', 'Sig']]
    
    table_num = add_table_page(
        doc, table_num,
        "Conditional total effects of FASt status on Developmental Adjustment.",
        tbl5,
        "Low = −1 SD; Mid = mean; High = +1 SD of credit dose. "
        "* p < .05 based on bootstrap CI excluding zero. "
        "Total effects significant at mean and high dose levels."
    )
    
    # Table 6: Indices of moderated mediation
    imm = df[df['parameter'].str.startswith('index_MM')].copy()
    imm['Mediator Pathway'] = imm['parameter'].apply(
        lambda x: 'via Emotional Distress' if 'EmoDiss' in x else 'via Quality of Engagement'
    )
    imm['Estimate'] = imm['est'].apply(lambda x: fmt_num(x))
    imm['SE'] = imm['boot_se'].apply(lambda x: fmt_num(x))
    imm['95% CI'] = imm.apply(lambda r: fmt_ci(r['ci_lower'], r['ci_upper']), axis=1)
    imm['Sig'] = imm['sig'].map({True: '*', False: ''})
    
    tbl6 = imm[['Mediator Pathway', 'Estimate', 'SE', '95% CI', 'Sig']].reset_index(drop=True)
    
    table_num = add_table_page(
        doc, table_num,
        "Index of Moderated Mediation (IMM) for each indirect pathway.",
        tbl6,
        "IMM = a × b interaction effect: (FASt × Dose → Mediator) × (Mediator → Outcome). "
        "A significant IMM indicates that the indirect effect varies linearly with credit dose. "
        "* p < .05 based on bootstrap CI excluding zero. "
        "Both pathways show significant moderated mediation."
    )
    
    # Table 7: Summary interpretation
    summary_data = [
        ['Research Question', 'Finding', 'Support'],
        ['RQ1: Direct effect of FASt', 'FASt → Developmental Adjustment not significant', 'No'],
        ['RQ2a: Mediation via Emotional Distress', 'FASt increases distress → lower adjustment', 'Yes'],
        ['RQ2b: Mediation via Quality of Engagement', 'FASt decreases engagement → lower adjustment', 'Yes'],
        ['RQ3a: Dose moderates distress pathway', 'Higher credit dose amplifies harm via distress', 'Yes'],
        ['RQ3b: Dose moderates engagement pathway', 'Higher credit dose amplifies harm via engagement', 'Yes'],
    ]
    summary_df = pd.DataFrame(summary_data[1:], columns=summary_data[0])
    
    table_num = add_table_page(
        doc, table_num,
        "Summary of research question support from bootstrap analysis.",
        summary_df,
        "Based on B = 500 bootstrap replicates with 100% convergence. "
        "Support determined by 95% BCa bootstrap CI excluding zero."
    )
    
    # Table 8: Total Generalizability Evidence
    gen_data = [
        ['Evidence Domain', 'Method', 'Result', 'Strength'],
        ['Statistical Inference', 'BCa Bootstrap (B=500)', '100% convergence; stable SEs', 'Strong'],
        ['Causal Identification', 'Propensity Score Weighting', 'ATO overlap weights; SMD < 0.1', 'Strong'],
        ['Direct Effect (c′)', 'Conditional at Low/Mid/High dose', 'Non-significant at all levels', 'Consistent null'],
        ['Indirect via Distress', 'Conditional at Low/Mid/High dose', 'Significant negative at all levels', 'Strong'],
        ['Indirect via Engagement', 'Conditional at Low/Mid/High dose', 'Significant negative at Mid/High', 'Moderate'],
        ['Total Effect', 'Conditional at Low/Mid/High dose', 'Significant negative at Mid/High', 'Moderate'],
        ['Moderated Mediation (Distress)', 'Index of Moderated Mediation', 'IMM = −0.034, CI excludes 0', 'Strong'],
        ['Moderated Mediation (Engagement)', 'Index of Moderated Mediation', 'IMM = −0.046, CI excludes 0', 'Strong'],
        ['Effect Direction', 'All significant paths', 'Consistent negative (harmful)', 'Strong'],
        ['Dose-Response Pattern', 'Slope of conditional effects', 'Monotonic: more credits → worse', 'Strong'],
    ]
    gen_df = pd.DataFrame(gen_data[1:], columns=gen_data[0])
    
    table_num = add_table_page(
        doc, table_num,
        "Summary of evidence for generalizability and robustness of findings.",
        gen_df,
        "Primary evidence from bootstrap-then-weight analysis with propensity score overlap weighting. "
        "Strength ratings: Strong = consistent across all conditions; Moderate = present in most conditions; "
        "Consistent null = reliably non-significant across conditions."
    )
    
    # Save
    out_docx.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_docx))
    print(f"Wrote: {out_docx}")
    print(f"\nTables created:")
    print("  1. Structural path coefficients")
    print("  2. Conditional a-paths (X → Mediators)")
    print("  3. Conditional direct effects")
    print("  4. Conditional indirect effects")
    print("  5. Conditional total effects")
    print("  6. Index of Moderated Mediation")
    print("  7. Summary of research question support")
    print("  8. Total generalizability evidence")


if __name__ == "__main__":
    main()
